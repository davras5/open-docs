#!/usr/bin/env python3
"""Generate all demo files for the OpenDocs construction project prototype."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers
from fpdf import FPDF
import csv

BASE = Path(r"c:/Users/DavidRasner/Documents/GitHub/open-docs/demo-files")

FOLDERS = [
    "01 Planung",
    "02 Genehmigungen",
    "03 Bauausfuehrung",
    "04 Kostenplanung",
    "05 Protokolle",
    "06 Fotos",
    "07 Vertraege",
    "08 Arbeitssicherheit",
]


def ensure_dirs():
    BASE.mkdir(parents=True, exist_ok=True)
    for f in FOLDERS:
        (BASE / f).mkdir(exist_ok=True)


# ── Helpers ──────────────────────────────────────────────────────────────

def make_xlsx(path, headers, rows, col_widths=None):
    """Create an xlsx with bold headers and sized columns."""
    wb = Workbook()
    ws = wb.active
    hfont = Font(bold=True, size=11)
    thin = Side(style="thin")
    border = Border(top=thin, left=thin, right=thin, bottom=thin)
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")

    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = hfont
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    for r, row in enumerate(rows, 2):
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.border = border
            cell.alignment = Alignment(wrap_text=True)

    if col_widths:
        for i, w in enumerate(col_widths):
            ws.column_dimensions[chr(65 + i) if i < 26 else "A" + chr(65 + i - 26)].width = w
    else:
        for c, h in enumerate(headers, 1):
            max_len = len(str(h))
            for row in rows:
                if c - 1 < len(row):
                    max_len = max(max_len, len(str(row[c - 1])))
            col_letter = chr(64 + c) if c <= 26 else "A" + chr(64 + c - 26)
            ws.column_dimensions[col_letter].width = min(max_len + 4, 40)

    wb.save(path)


def safe(text):
    """Encode text to latin-1 for fpdf, replacing unsupported chars."""
    return text.encode("latin-1", "replace").decode("latin-1")


class GermanPDF(FPDF):
    """PDF with latin-1 safe text output."""

    def s_cell(self, w, h=10, txt="", border=0, align="", new_line=False, fill=False):
        """Cell that safely encodes text. new_line=True moves to next line after."""
        from fpdf.enums import XPos, YPos
        if new_line:
            self.cell(w, h, safe(txt), border=border, align=align, fill=fill,
                      new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        else:
            self.cell(w, h, safe(txt), border=border, align=align, fill=fill,
                      new_x=XPos.RIGHT, new_y=YPos.TOP)

    def s_multi_cell(self, w, h, txt, border=0, align="J", fill=False):
        from fpdf.enums import XPos, YPos
        self.multi_cell(w, h, safe(txt), border=border, align=align, fill=fill,
                        new_x=XPos.LMARGIN, new_y=YPos.LAST)

    def header_block(self, title, subtitle=""):
        self.set_font("Helvetica", "B", 16)
        self.s_cell(0, 12, title, align="C", new_line=True)
        if subtitle:
            self.set_font("Helvetica", "", 11)
            self.s_cell(0, 8, subtitle, align="C", new_line=True)
        self.ln(5)

    def section(self, title):
        self.set_font("Helvetica", "B", 12)
        self.s_cell(0, 10, title, new_line=True)
        self.set_font("Helvetica", "", 10)

    def body(self, text):
        self.set_font("Helvetica", "", 10)
        self.s_multi_cell(0, 5, text)
        self.ln(2)

    def bullet(self, text):
        self.set_font("Helvetica", "", 10)
        self.s_multi_cell(0, 5, "  - " + text)


def add_docx_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Table Grid")
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            table.rows[r + 1].cells[c].text = str(val)
    return table


# ── ROOT FILES ───────────────────────────────────────────────────────────

def create_projektuebersicht_md():
    content = """# Neubau Bürogebäude — TechPark München Ost

## Projektdaten

| Eigenschaft | Wert |
|-------------|------|
| Bauherr | Müller & Schmidt Immobilien GmbH |
| Standort | Messestraße 42, 81829 München |
| Generalplaner | Architekten Lechner & Kollegen PartGmbB |
| Bauleitung | Berger Baumanagement GmbH |
| Baubeginn | 01.03.2026 |
| Fertigstellung | 15.09.2028 (geplant) |
| Gesamtbudget | 38,5 Mio. EUR (brutto) |

## Ordnerstruktur

- 01 Planung — Projektbeschreibung, Raumprogramm, Terminplan
- 02 Genehmigungen — Baugenehmigung, Gutachten, Brandschutz
- 03 Bauausführung — Baustellenordnung, Nachunternehmer, Tagesberichte
- 04 Kostenplanung — DIN 276, Zahlungspläne, Nachträge
- 05 Protokolle — Baubesprechungen, Abnahmen
- 06 Fotos — Baudokumentation
- 07 Verträge — Planervertrag, Bürgschaften
- 08 Arbeitssicherheit — SiGePlan, Unterweisungen

## Ansprechpartner

| Name | Rolle | Telefon |
|------|-------|---------|
| Dr. Martin Müller | Bauherr | 089 / 12345-100 |
| Arch. Sabine Lechner | Generalplanerin | 089 / 98765-200 |
| Dipl.-Ing. Thomas Berger | Bauleiter | 0171 555 2341 |
| Ing. Andrea Hofmann | SiGeKo | 0172 888 4567 |
| Michael Gruber | Polier | 0170 333 9876 |
"""
    (BASE / "Projektuebersicht.md").write_text(content, encoding="utf-8")


def create_kontaktliste_xlsx():
    headers = ["Name", "Firma", "Rolle", "E-Mail", "Telefon", "Mobiltelefon"]
    rows = [
        ["Dr. Martin Müller", "Müller & Schmidt Immobilien GmbH", "Bauherr", "mueller@ms-immo.de", "089 12345-100", "0171 111 2233"],
        ["Arch. Sabine Lechner", "Architekten Lechner & Kollegen", "Generalplanerin", "lechner@arch-lk.de", "089 98765-200", "0172 222 3344"],
        ["Dipl.-Ing. Thomas Berger", "Berger Baumanagement GmbH", "Bauleiter", "berger@bau-mg.de", "089 54321-50", "0171 555 2341"],
        ["Ing. Andrea Hofmann", "Hofmann Sicherheitstechnik", "SiGeKo", "hofmann@sigeko.de", "089 77777-10", "0172 888 4567"],
        ["Hr. Josef Bauer", "Bauer Tiefbau GmbH", "Erdarbeiten", "bauer@bauer-tiefbau.de", "089 66666-30", "0170 444 5566"],
        ["Fr. Claudia Kellner", "Züblin AG, NL München", "Rohbau", "kellner@zueblin.de", "089 55555-20", "0171 666 7788"],
        ["Michael Gruber", "Bauer Tiefbau GmbH", "Polier", "gruber@bauer-tiefbau.de", "-", "0170 333 9876"],
        ["Hr. Werner Weiss", "Metallbau Schüco Partner", "Fassade", "weiss@schueco-partner.de", "089 44444-15", "0172 999 0011"],
        ["Hr. Karl Müller", "Elektro Müller GmbH", "Elektro", "k.mueller@elektro-mueller.de", "089 33333-25", "0173 111 2244"],
        ["Fr. Petra Wagner", "Haustechnik Süd GmbH", "HLKS", "wagner@ht-sued.de", "089 22222-35", "0174 222 3355"],
    ]
    make_xlsx(BASE / "Kontaktliste.xlsx", headers, rows, [28, 35, 18, 30, 18, 18])


# ── 01 PLANUNG ───────────────────────────────────────────────────────────

def create_projektbeschreibung_docx():
    doc = Document()
    doc.add_heading("Projektbeschreibung", level=1)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("Projekt: ").bold = True
    p.add_run('Neubau Bürogebäude "TechPark München Ost"')
    p = doc.add_paragraph()
    p.add_run("Bauherr: ").bold = True
    p.add_run("Müller & Schmidt Immobilien GmbH")
    p = doc.add_paragraph()
    p.add_run("Standort: ").bold = True
    p.add_run("Messestraße 42, 81829 München")
    p = doc.add_paragraph()
    p.add_run("Generalplaner: ").bold = True
    p.add_run("Architekten Lechner & Kollegen PartGmbB")

    doc.add_heading("Eckdaten", level=2)
    add_docx_table(doc, ["Eigenschaft", "Wert"], [
        ["Grundstücksfläche", "4.200 m²"],
        ["Bruttogrundfläche (BGF)", "12.500 m²"],
        ["Geschosse", "6 Obergeschosse + 2 Untergeschosse"],
        ["Stellplätze", "180 (TG) + 40 (Außen)"],
        ["Gebäudehöhe", "ca. 24,5 m"],
    ])

    doc.add_heading("Geplante Nutzung", level=2)
    usages = [
        "EG: Empfang, Lobby, Konferenzbereich, Café (ca. 1.800 m²)",
        "1.–4. OG: Open-Space-Büroflächen, Besprechungsräume (je ca. 1.500 m²)",
        "5. OG: Geschäftsleitung, Executive Lounge, Dachterrasse (ca. 1.200 m²)",
        "UG1: Tiefgarage (120 Stellplätze), Fahrradstellplätze (ca. 2.100 m²)",
        "UG2: Tiefgarage (60 Stellplätze), Haustechnik, Lager (ca. 2.100 m²)",
    ]
    for u in usages:
        doc.add_paragraph(u, style="List Bullet")

    doc.add_heading("Bauweise und Nachhaltigkeit", level=2)
    doc.add_paragraph(
        "Das Gebäude wird in Stahlbetonkonstruktion mit einer vorgehängten "
        "Glasfassade errichtet. Die Planung orientiert sich am KfW-Effizienzhaus-40-Standard "
        "und strebt eine DGNB-Zertifizierung in Gold an."
    )
    items = [
        "Stahlbetonkonstruktion mit Flachdecken",
        "Elementierte Glasfassade (Dreifachverglasung, Sonnenschutz)",
        "Bauteilaktivierung in den Geschossdecken",
        "Photovoltaikanlage auf dem Dach (ca. 350 kWp)",
        "Gründach mit Retentionsfunktion",
        "Regenwassernutzung für Sanitäranlagen und Bewässerung",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Zeitraum und Budget", level=2)
    doc.add_paragraph(
        "Der Bauablauf ist von März 2026 bis September 2028 geplant (ca. 30 Monate). "
        "Das Gesamtbudget beträgt ca. 38,5 Mio. EUR brutto inkl. Planungskosten, "
        "Außenanlagen und Reserven."
    )

    doc.save(str(BASE / "01 Planung" / "Projektbeschreibung.docx"))


def create_raumprogramm_xlsx():
    headers = ["Geschoss", "Bereich", "Fläche (m²)", "Anzahl Arbeitsplätze", "Bemerkung"]
    rows = [
        ["EG", "Empfang / Lobby", 350, "-", "Repräsentativ, doppelte Raumhöhe"],
        ["EG", "Konferenzbereich", 500, "-", "4 Räume à 20–50 Plätze"],
        ["EG", "Café / Kantine", 300, "-", "Inkl. Küche"],
        ["EG", "Technik / Nebenräume", 250, "-", ""],
        ["EG", "Fahrradraum / Umkleiden", 200, "-", "Duschen vorhanden"],
        ["1. OG", "Open-Space-Büro", 1100, 110, "Flexible Möblierung"],
        ["1. OG", "Besprechungsräume", 250, "-", "6 Räume"],
        ["1. OG", "Teeküche / Sozialräume", 150, "-", ""],
        ["2. OG", "Open-Space-Büro", 1100, 110, "Flexible Möblierung"],
        ["2. OG", "Besprechungsräume", 250, "-", "6 Räume"],
        ["2. OG", "Teeküche / Sozialräume", 150, "-", ""],
        ["3. OG", "Open-Space-Büro", 1100, 100, ""],
        ["3. OG", "Besprechungsräume", 250, "-", "4 Räume"],
        ["3. OG", "Serverraum", 50, "-", "Klimatisiert"],
        ["4. OG", "Open-Space-Büro", 1100, 100, ""],
        ["4. OG", "Besprechungsräume", 250, "-", "4 Räume"],
        ["5. OG", "Geschäftsleitung", 600, 20, "Einzelbüros + Assistenz"],
        ["5. OG", "Executive Lounge", 300, "-", "Inkl. Dachterrasse"],
        ["UG1", "Tiefgarage", 2100, "-", "120 Stellplätze"],
        ["UG2", "Tiefgarage + Technik", 2100, "-", "60 Stellplätze + Haustechnik"],
    ]
    make_xlsx(BASE / "01 Planung" / "Raumprogramm.xlsx", headers, rows, [12, 28, 14, 22, 32])


def create_terminplan_pdf():
    from fpdf.enums import XPos, YPos
    pdf = GermanPDF()
    pdf.add_page()
    pdf.header_block("Terminplan", safe("Neubau Bürogebäude München — TechPark München Ost"))

    pdf.set_font("Helvetica", "", 9)
    pdf.s_cell(0, 6, "Stand: 01.03.2026 | Erstellt: Berger Baumanagement GmbH", align="C", new_line=True)
    pdf.ln(5)

    pdf.section("Meilensteine")
    pdf.ln(2)

    milestones = [
        ["M1", "Baubeginn / Baustelleneinrichtung", "01.03.2026", "Abgeschlossen"],
        ["M2", "Beginn Erdarbeiten / Spundwand", "15.03.2026", "In Arbeit"],
        ["M3", "Baugrube fertiggestellt", "30.06.2026", "Offen"],
        ["M4", "Rohbau UG abgeschlossen", "30.09.2026", "Offen"],
        ["M5", safe("Rohbau EG-3. OG"), "28.02.2027", "Offen"],
        ["M6", "Rohbau komplett (inkl. Dach)", "30.06.2027", "Offen"],
        ["M7", "Fassade geschlossen", "31.10.2027", "Offen"],
        ["M8", "Beginn Innenausbau", "01.09.2027", "Offen"],
        ["M9", "Technische Inbetriebnahme", "30.06.2028", "Offen"],
        ["M10", "Fertigstellung / Abnahme", "15.09.2028", "Offen"],
    ]

    col_w = [15, 80, 35, 30]
    pdf.set_font("Helvetica", "B", 9)
    for w, h in zip(col_w, ["Nr.", "Meilenstein", "Termin", "Status"]):
        pdf.cell(w, 7, safe(h), border=1, align="C", new_x=XPos.RIGHT, new_y=YPos.TOP)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    for row in milestones:
        for w, val in zip(col_w, row):
            pdf.cell(w, 6, safe(val), border=1, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.ln()

    pdf.ln(8)
    pdf.section("Kritischer Pfad")
    pdf.body(
        safe("Der kritische Pfad verlaeuft ueber Erdarbeiten -> Rohbau UG -> Rohbau OG -> "
        "Fassade -> Innenausbau -> Technische Inbetriebnahme. Eine Verzoegerung der "
        "Spundwandarbeiten um mehr als 2 Wochen gefaehrdet den Meilenstein M3.")
    )

    pdf.section("Hinweise")
    pdf.bullet(safe("Winterpause: Betonarbeiten ab Dezember witterungsabhaengig"))
    pdf.bullet(safe("Fassadenelemente: Lieferzeit ca. 16 Wochen, fruehzeitige Bestellung erforderlich"))
    pdf.bullet(safe("Koordination TGA mit Rohbau ab M5 engmaschig planen"))

    pdf.output(str(BASE / "01 Planung" / "Terminplan_Uebersicht.pdf"))


def create_architekturbriefing_docx():
    doc = Document()
    doc.add_heading("Architekturbriefing", level=1)
    doc.add_paragraph(
        "Grundlage für die architektonische Gestaltung des Neubaus "
        '"TechPark München Ost", Messestraße 42, 81829 München.'
    )

    doc.add_heading("Gestaltungskonzept", level=2)
    doc.add_paragraph(
        "Das Gebäude soll als modernes, transparentes Bürogebäude in Erscheinung treten "
        "und sich harmonisch in den entstehenden TechPark einfügen. Leitgedanke ist eine "
        "offene, kommunikationsfördernde Architektur mit hoher Aufenthaltsqualität."
    )
    for item in [
        "Klare Kubatur mit horizontaler Gliederung",
        "Großzügige Glasflächen für maximalen Tageslichteinfall",
        "Rücksprung im 5. OG mit begrünter Dachterrasse",
        "Repräsentativer Eingangsbereich mit doppelter Raumhöhe",
        "Sichtbare Holzelemente im Innenbereich als Kontrast zu Glas und Beton",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Materialien", level=2)
    add_docx_table(doc, ["Bauteil", "Material", "Bemerkung"], [
        ["Tragwerk", "Stahlbeton C30/37", "Flachdeckensystem"],
        ["Fassade", "Aluminium-Glas-Elementfassade", "Dreifachverglasung, Uw ≤ 0,9"],
        ["Sonnenschutz", "Außenraffstore, elektrisch", "Windwächter-gesteuert"],
        ["Bodenbeläge Büro", "Doppelboden mit Teppichfliesen", "Hohlraumhöhe 15 cm"],
        ["Bodenbeläge Foyer", "Naturstein (Jura-Kalkstein)", "Rutschklasse R10"],
        ["Innenwände", "Trockenbau / Systemtrennwände", "Flexible Raumaufteilung"],
        ["Dach", "Flachdach, extensiv begrünt", "Retentionsdach mit PV-Aufständerung"],
    ])

    doc.add_heading("Nachhaltigkeitsanforderungen", level=2)
    for item in [
        "KfW-Effizienzhaus 40 (Primärenergiebedarf ≤ 40 % des Referenzgebäudes)",
        "DGNB-Zertifizierung in Gold angestrebt",
        "Ökobilanz (LCA) für die Tragkonstruktion erforderlich",
        "Recyclingfähige Materialien bevorzugen (Cradle-to-Cradle wo möglich)",
        "Regenwassermanagement: Retentionsdach + Zisterne für Grauwassernutzung",
        "Barrierefreiheit nach DIN 18040-1",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(BASE / "01 Planung" / "Architekturbriefing.docx"))


# ── 02 GENEHMIGUNGEN ────────────────────────────────────────────────────

def create_baugenehmigung_pdf():
    pdf = GermanPDF()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 14)
    pdf.s_cell(0, 10, "LANDESHAUPTSTADT MÜNCHEN", new_line=True, align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.s_cell(0, 7, "Referat für Stadtplanung und Bauordnung", new_line=True, align="C")
    pdf.ln(3)
    pdf.set_draw_color(0, 0, 0)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(8)

    pdf.set_font("Helvetica", "B", 16)
    pdf.s_cell(0, 12, "BAUGENEHMIGUNGSBESCHEID", new_line=True, align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "", 10)
    pdf.s_cell(40, 6, "Aktenzeichen:")
    pdf.set_font("Helvetica", "B", 10)
    pdf.s_cell(0, 6, "LBK-2025-48291-BG", new_line=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.s_cell(40, 6, "Datum:")
    pdf.s_cell(0, 6, "15.02.2026", new_line=True)
    pdf.s_cell(40, 6, "Bauherr:")
    pdf.s_cell(0, 6, safe("Müller & Schmidt Immobilien GmbH"), new_line=True)
    pdf.s_cell(40, 6, "Baugrundstück:")  # Changed from Baugrundstück
    pdf.s_cell(0, 6, safe("Messestraße 42, 81829 München"), new_line=True)
    pdf.s_cell(40, 6, "Vorhaben:")
    pdf.s_cell(0, 6, safe("Neubau Bürogebäude (6 OG + 2 UG)"), new_line=True)
    pdf.ln(5)

    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(5)

    pdf.section("I. Genehmigung")
    pdf.body(
        "Auf Grundlage der eingereichten Bauvorlagen wird hiermit die Baugenehmigung "
        "für den Neubau eines sechsgeschossigen Bürogebäudes mit zwei Untergeschossen "
        "(Tiefgarage) auf dem Grundstück Fl.-Nr. 1234/56, Gemarkung Trudering, erteilt. "
        "Das Vorhaben entspricht den Festsetzungen des Bebauungsplans Nr. 2087a."
    )

    pdf.section("II. Auflagen")
    auflagen = [
        "Die Vorgaben des genehmigten Brandschutzkonzepts (Anlage 5) sind vollständig umzusetzen.",
        "Während der Bauphase ist ein qualifizierter Sicherheits- und Gesundheitsschutzkoordinator (SiGeKo) zu bestellen.",
        "Die Erdarbeiten sind durch einen Bodengutachter zu begleiten; Aushubmaterial ist ordnungsgemäß zu entsorgen (Entsorgungsnachweis erforderlich).",
        "Die Immissionsrichtwerte der AVV Baulärm sind einzuhalten. Lärmintensive Arbeiten nur Mo–Fr 07:00–20:00 Uhr, Sa 08:00–16:00 Uhr.",
        "Vor Baubeginn ist eine Beweissicherung der angrenzenden Bestandsgebäude durchzuführen und zu dokumentieren.",
    ]
    for i, a in enumerate(auflagen, 1):
        pdf.set_font("Helvetica", "B", 10)
        pdf.s_cell(8, 5, f"{i}.")
        pdf.set_font("Helvetica", "", 10)
        pdf.s_multi_cell(0, 5, safe(a))
        pdf.ln(2)

    pdf.ln(3)
    pdf.section("III. Rechtsbehelfsbelehrung")
    pdf.body(
        "Gegen diesen Bescheid kann innerhalb eines Monats nach Bekanntgabe "
        "Widerspruch bei der Landeshauptstadt München, Referat für Stadtplanung "
        "und Bauordnung, erhoben werden."
    )

    pdf.ln(10)
    pdf.set_font("Helvetica", "", 10)
    pdf.s_cell(0, 6, safe("München, den 15.02.2026"), new_line=True)
    pdf.ln(15)
    pdf.s_cell(0, 6, "Dr. Elisabeth Brandner", new_line=True)
    pdf.s_cell(0, 6, safe("Abteilungsleiterin Bauordnung"), new_line=True)

    pdf.output(str(BASE / "02 Genehmigungen" / "Baugenehmigung_Bescheid.pdf"))


def create_umweltvertraeglichkeit_docx():
    doc = Document()
    doc.add_heading("Umweltverträglichkeitsprüfung", level=1)
    doc.add_paragraph(
        "Prüfung der Umweltauswirkungen für das Vorhaben Neubau Bürogebäude "
        '"TechPark München Ost", Messestraße 42, 81829 München.'
    )

    for title, text in [
        ("Schutzgut Boden", "Das Baugrundstück war vormals gewerblich genutzt (ehem. Lagerfläche). "
         "Ein Bodengutachten (Anlage B1) bestätigt, dass keine Altlasten vorliegen. "
         "Der Oberboden (ca. 2.500 m³) wird getrennt abgetragen und fachgerecht zwischengelagert. "
         "Die Versiegelung nimmt durch die geplante Tiefgarage zu, wird aber durch Dachbegrünung "
         "und Entsiegelung im Umfeld teilweise kompensiert."),
        ("Schutzgut Wasser", "Die Grundwasserflurabstände liegen bei ca. 4,5 m unter GOK. "
         "Die Baugrube reicht bis ca. 7 m Tiefe, weshalb eine temporäre Wasserhaltung "
         "erforderlich wird. Das anfallende Grundwasser wird nach Aufbereitung in den "
         "Vorfluter (Hachinger Bach) eingeleitet (wasserrechtliche Erlaubnis beantragt). "
         "Regenwasser wird über Retentionsdächer und eine Zisterne bewirtschaftet."),
        ("Schutzgut Luft und Klima", "Während der Bauphase ist mit erhöhter Staubentwicklung "
         "bei Erdarbeiten zu rechnen. Maßnahmen: Befeuchtung der Fahrwege, Abdeckung "
         "von Transportfahrzeugen. Im Betrieb wird das Gebäude emissionsfrei beheizt "
         "(Wärmepumpe + Fernwärme). Die Dachbegrünung und Baumpflanzungen im Außenbereich "
         "verbessern das Mikroklima."),
        ("Schutzgut Flora und Fauna", "Auf dem Grundstück befinden sich keine geschützten "
         "Arten oder Biotope. Fünf bestehende Bäume (Bergahorn, Birke) müssen gefällt werden; "
         "als Ausgleich werden 12 Neupflanzungen (heimische Laubbäume) im Außenbereich vorgesehen. "
         "Artenschutzfachliche Untersuchung gemäß saP liegt vor (Anlage F1)."),
    ]:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)

    doc.add_heading("Zusammenfassung und Bewertung", level=2)
    doc.add_paragraph(
        "Die Umweltverträglichkeitsprüfung ergibt, dass bei Einhaltung der genannten "
        "Vermeidungs- und Minimierungsmaßnahmen keine erheblichen nachteiligen "
        "Umweltauswirkungen zu erwarten sind. Das Vorhaben ist aus umweltfachlicher "
        "Sicht genehmigungsfähig."
    )

    doc.save(str(BASE / "02 Genehmigungen" / "Umweltvertraeglichkeitspruefung.docx"))


def create_brandschutzkonzept_docx():
    doc = Document()
    doc.add_heading("Brandschutzkonzept — Entwurf", level=1)
    doc.add_paragraph(
        "Bürogebäude TechPark München Ost, Messestraße 42, 81829 München. "
        "Erstellt gemäß BayBO und Richtlinien der vfdb."
    )

    doc.add_heading("Gebäudeklasse und Einstufung", level=2)
    doc.add_paragraph(
        "Gebäudeklasse 5 gemäß Art. 2 BayBO (Gebäudehöhe ca. 24,5 m, "
        "mehr als 1.600 m² Geschossfläche). Es handelt sich um einen Sonderbau "
        "gemäß Art. 2 Abs. 4 BayBO (Bürogebäude mit mehr als 400 m² Nutzfläche pro Geschoss)."
    )

    doc.add_heading("Rettungswege", level=2)
    for item in [
        "Zwei bauliche Rettungswege pro Geschoss (Treppenräume TR1 und TR2)",
        "Treppenräume als notwendige Treppenräume mit Druckbelüftung",
        "Maximale Rettungsweglänge: 35 m (eingehalten in allen Geschossen)",
        "Barrierefreie Rettung über Aufzug mit Feuerwehr-Bedienung (Evakuierungsaufzug)",
        "Fluchtwegbeschilderung gemäß DIN ISO 7010, nachleuchtend",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Brandabschnitte", level=2)
    doc.add_paragraph(
        "Das Gebäude wird in zwei Brandabschnitte unterteilt (Achse C als Brandwand F90). "
        "Die Untergeschosse bilden jeweils eigene Brandabschnitte. "
        "Alle Durchbrüche in Brandwänden sind mit zugelassenen Abschottungen zu versehen."
    )

    doc.add_heading("Löschanlagen", level=2)
    for item in [
        "Trockene Steigleitung in beiden Treppenräumen",
        "Wandhydranten Typ F auf jeder Ebene",
        "Sprinkleranlage in UG1 und UG2 (Tiefgarage) gemäß VdS CEA 4001",
        "Aufschaltung auf Feuerwehr-Gebäudefunkanlage",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Rauchableitung", level=2)
    doc.add_paragraph(
        "Natürliche Rauch- und Wärmeabzugsanlagen (NRWG) im Dachbereich "
        "mit einer aerodynamisch wirksamen Öffnungsfläche von mind. 1,5 % "
        "der Grundfläche. In den UG-Geschossen maschinelle Entrauchung "
        "über Entrauchungsschächte."
    )

    doc.add_heading("Brandmeldeanlage", level=2)
    doc.add_paragraph(
        "Vollflächige Brandmeldeanlage (BMA) der Kategorie 1 gemäß DIN 14675. "
        "Aufschaltung auf die Integrierte Leitstelle München. "
        "Feuerwehr-Informations- und Bediensystem (FIBS) im Eingangsbereich."
    )

    doc.save(str(BASE / "02 Genehmigungen" / "Brandschutzkonzept_Entwurf.docx"))


# ── 03 BAUAUSFÜHRUNG ─────────────────────────────────────────────────────

def create_baustellenordnung_pdf():
    pdf = GermanPDF()
    pdf.add_page()
    pdf.header_block("Baustellenordnung", "Neubau TechPark München Ost — Messestraße 42")
    pdf.set_font("Helvetica", "", 9)
    pdf.s_cell(0, 6, "Gültig ab: 01.03.2026 | Version 1.0", new_line=True, align="C")
    pdf.ln(5)

    sections = [
        ("1. Arbeitszeiten", [
            "Montag bis Freitag: 07:00 – 18:00 Uhr",
            "Samstag (nach Genehmigung): 08:00 – 16:00 Uhr",
            "Sonn- und Feiertage: Kein Baubetrieb",
            "Lärmintensive Arbeiten: Mo–Fr 07:00–20:00, Sa 08:00–16:00 (gemäß AVV Baulärm)",
        ]),
        ("2. Zugang und Zufahrt", [
            "Baustellenzufahrt ausschließlich über Tor 1 (Messestraße, Nordseite)",
            "Alle Personen müssen sich am Baucontainer anmelden und den Baustellenausweis sichtbar tragen",
            "LKW-Verkehr: Max. 7,5 t über Tor 1; Schwertransporte nur nach Voranmeldung über Tor 2",
            "Geschwindigkeitsbegrenzung auf dem Baufeld: 10 km/h",
        ]),
        ("3. Ordnung und Sauberkeit", [
            "Jeder Nachunternehmer ist für die Sauberkeit seines Arbeitsbereichs verantwortlich",
            "Abfälle sind nach Fraktionen getrennt in den dafür vorgesehenen Containern zu entsorgen",
            "Verschmutzung öffentlicher Straßen ist sofort zu beseitigen (Reifenwaschanlage nutzen)",
            "Materialien sind ordentlich zu lagern; Fluchtwege stets freizuhalten",
        ]),
        ("4. Sicherheit", [
            "PSA-Pflicht: Helm, Sicherheitsschuhe S3, Warnweste auf dem gesamten Baufeld",
            "Zusätzliche PSA je nach Tätigkeit (Gehörschutz, Schutzbrille, Absturzsicherung)",
            "Rauchen nur in der ausgewiesenen Raucherzone (Container 7)",
            "Alkohol- und Drogenverbot auf der gesamten Baustelle",
            "Ersthelfer: Michael Gruber (Polier) — Tel. 0170 333 9876",
            "Sammelplatz bei Evakuierung: Parkplatz Nordost (gekennzeichnet)",
        ]),
        ("5. Ansprechpartner", [
            "Bauleitung: Dipl.-Ing. Thomas Berger — 0171 555 2341",
            "Polier: Michael Gruber — 0170 333 9876",
            "SiGeKo: Ing. Andrea Hofmann — 0172 888 4567",
            "Bauherrenvertreter: Dr. Martin Müller — 089 12345-100",
        ]),
    ]

    for title, items in sections:
        pdf.section(title)
        for item in items:
            pdf.bullet(item)
        pdf.ln(4)

    pdf.output(str(BASE / "03 Bauausfuehrung" / "Baustellenordnung.pdf"))


def create_nachunternehmer_xlsx():
    headers = ["Gewerk", "Firma", "Ansprechpartner", "Vertragssumme (EUR)", "Status", "Beginn", "Ende"]
    rows = [
        ["Erdarbeiten / Spezialtiefbau", "Bauer Tiefbau GmbH", "Hr. Josef Bauer", 2850000, "Beauftragt", "15.03.2026", "30.06.2026"],
        ["Rohbau", "Züblin AG, NL München", "Fr. Claudia Kellner", 8200000, "Beauftragt", "01.07.2026", "30.06.2027"],
        ["Fassade", "Metallbau Schüco Partner", "Hr. Werner Weiss", 3400000, "Beauftragt", "01.07.2027", "31.10.2027"],
        ["Elektroinstallation", "Elektro Müller GmbH", "Hr. Karl Müller", 2100000, "Beauftragt", "01.09.2027", "30.06.2028"],
        ["HLKS (Heizung/Lüftung/Klima/Sanitär)", "Haustechnik Süd GmbH", "Fr. Petra Wagner", 3800000, "Beauftragt", "01.09.2027", "30.06.2028"],
        ["Aufzüge", "Schindler Deutschland AG", "Hr. Markus Stein", 680000, "In Verhandlung", "01.01.2028", "31.05.2028"],
        ["Trockenbau / Innenausbau", "Knauf Interiors GmbH", "Fr. Lisa Maier", 1900000, "In Verhandlung", "01.11.2027", "30.04.2028"],
        ["Bodenbeläge", "Carpet Concept / Jura Marmor", "Hr. Peter Hofmann", 750000, "Angebot eingeholt", "01.02.2028", "31.05.2028"],
        ["Malerarbeiten", "Malerfachbetrieb Riedl", "Hr. Georg Riedl", 420000, "Angebot eingeholt", "01.03.2028", "31.07.2028"],
        ["Außenanlagen / Landschaftsbau", "Garten- und Landschaftsbau Meier", "Hr. Stefan Meier", 580000, "Noch nicht vergeben", "01.05.2028", "31.08.2028"],
        ["Sprinkler / Brandschutz", "Minimax GmbH", "Hr. Andreas Koch", 450000, "Beauftragt", "01.10.2027", "28.02.2028"],
    ]
    make_xlsx(BASE / "03 Bauausfuehrung" / "Nachunternehmer_Uebersicht.xlsx", headers, rows,
              [34, 30, 22, 20, 18, 14, 14])


def create_bautagesbericht_docx():
    doc = Document()
    doc.add_heading("Bautagesbericht", level=1)

    add_docx_table(doc, ["Feld", "Angabe"], [
        ["Projekt", "Neubau TechPark München Ost"],
        ["Datum", "15.03.2026 (Montag)"],
        ["Bericht-Nr.", "BT-2026-011"],
        ["Erstellt von", "Dipl.-Ing. Thomas Berger"],
        ["Wetter", "Bewölkt, 8 °C, trocken, Wind SW 15 km/h"],
    ])
    doc.add_paragraph("")

    doc.add_heading("Personal auf der Baustelle", level=2)
    add_docx_table(doc, ["Firma", "Gewerk", "Anzahl Personen"], [
        ["Bauer Tiefbau GmbH", "Erdarbeiten / Spundwand", "12"],
        ["Berger Baumanagement", "Bauleitung / Überwachung", "3"],
        ["Vermessungsbüro Krämer", "Vermessung", "2"],
        ["Hofmann Sicherheitstechnik", "SiGeKo-Begehung", "1"],
    ])

    doc.add_heading("Geräte und Maschinen", level=2)
    for item in [
        "1x Seilbagger (Liebherr HS 8100) — Spundwandarbeiten",
        "1x Bagger (Cat 320) — Erdaushub",
        "2x LKW (MAN TGS) — Bodentransport",
        "1x Rüttelplatte — Verdichtungsarbeiten",
        "1x Radlader (Liebherr L 550) — Materialumschlag",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Ausgeführte Arbeiten", level=2)
    for item in [
        "Fortführung der Spundwandarbeiten Achse 1–3 (Nordseite), 14 Bohlen eingebracht",
        "Aushub Baugrubenbecken Bereich A, ca. 280 m³ Boden abgefahren (Deponie Feldkirchen)",
        "Vermessung der Spundwandlage — Abweichung innerhalb Toleranz (±2 cm)",
        "Einrichtung Baustraße Abschnitt 2 abgeschlossen",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Besondere Vorkommnisse", level=2)
    doc.add_paragraph(
        "Keine besonderen Vorkommnisse. SiGeKo-Begehung ohne Beanstandungen."
    )

    doc.add_heading("Geplante Arbeiten (nächster Tag)", level=2)
    for item in [
        "Weiterführung Spundwand Achse 3–5",
        "Beginn Aushub Bereich B",
        "Anlieferung weiterer Spundbohlen (Lieferant: ArcelorMittal)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(BASE / "03 Bauausfuehrung" / "Bautagesbericht_2026-03-15.docx"))


def create_maengelliste_xlsx():
    headers = ["Nr.", "Gewerk", "Beschreibung", "Ort", "Festgestellt am", "Frist", "Status", "Verantwortlich"]
    rows = [
        ["M-001", "Erdarbeiten", "Verdichtung Planum Bereich A unter Sollwert (Ev2 = 38 MN/m², Soll ≥ 45 MN/m²)", "Baugrube Nord, Achse A1-A3", "18.03.2026", "25.03.2026", "Offen", "Bauer Tiefbau GmbH"],
        ["M-002", "Baustelleneinrichtung", "Bauzaunfeld Nr. 14 beschädigt, kippt bei Wind", "Westseite, Höhe Container 5", "20.03.2026", "22.03.2026", "Behoben", "Berger Baumanagement"],
        ["M-003", "Spezialtiefbau", "Spundbohle Nr. 27 ca. 8 cm aus Flucht — Nachvermessung erforderlich", "Nordseite, Achse 2", "22.03.2026", "29.03.2026", "In Prüfung", "Bauer Tiefbau GmbH"],
    ]
    make_xlsx(BASE / "03 Bauausfuehrung" / "Maengelliste.xlsx", headers, rows,
              [8, 18, 50, 28, 16, 14, 14, 22])


# ── 04 KOSTENPLANUNG ────────────────────────────────────────────────────

def create_kostenschaetzung_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "DIN 276 Kostenschätzung"

    hfont = Font(bold=True, size=11)
    thin = Side(style="thin")
    border = Border(top=thin, left=thin, right=thin, bottom=thin)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, size=11, color="FFFFFF")
    group_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    total_fill = PatternFill(start_color="FFC000", end_color="FFC000", fill_type="solid")

    headers = ["KG", "Kostengruppe", "Bezeichnung", "Kosten netto (EUR)", "Kosten brutto (EUR)", "Anteil (%)"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    data = [
        # (KG, group, name, netto, brutto, pct, is_group_header)
        (100, 1, "Grundstück", 3200000, 3808000, 9.9, True),
        (200, 1, "Vorbereitende Maßnahmen", 850000, 1011500, 2.6, True),
        (300, 1, "Bauwerk — Baukonstruktionen", None, None, None, True),
        (310, 2, "Baugrube / Gründung", 2400000, 2856000, 7.4, False),
        (320, 2, "Rohbau (Wände, Decken, Treppen)", 5200000, 6188000, 16.1, False),
        (330, 2, "Fassade / Außenwände", 2800000, 3332000, 8.7, False),
        (340, 2, "Dach", 450000, 535500, 1.4, False),
        ("300*", 1, "Summe KG 300", 10850000, 12911500, 33.6, True),
        (400, 1, "Bauwerk — Technische Anlagen", None, None, None, True),
        (410, 2, "Heizung / Kälte", 1800000, 2142000, 5.6, False),
        (420, 2, "Lüftung / Klimatisierung", 1500000, 1785000, 4.6, False),
        (430, 2, "Sanitär", 650000, 773500, 2.0, False),
        (440, 2, "Elektroinstallation", 2100000, 2499000, 6.5, False),
        (450, 2, "Aufzüge / Fördertechnik", 680000, 809200, 2.1, False),
        (460, 2, "Brandschutz / Sprinkler", 450000, 535500, 1.4, False),
        ("400*", 1, "Summe KG 400", 7180000, 8543200, 22.2, True),
        (500, 1, "Außenanlagen / Freiflächen", 1200000, 1428000, 3.7, True),
        (600, 1, "Ausstattung / Kunstwerke", 350000, 416500, 1.1, True),
        (700, 1, "Baunebenkosten (Planung, Gutachten)", 3700000, 4403000, 11.4, True),
        ("Rsv", 1, "Reserve / Unvorhergesehenes (5 %)", 1600000, 1904000, 5.0, True),
    ]

    row_idx = 2
    for kg, level, name, netto, brutto, pct, is_group in data:
        ws.cell(row=row_idx, column=1, value=str(kg)).border = border
        ws.cell(row=row_idx, column=2, value="" if level == 2 else "").border = border
        ws.cell(row=row_idx, column=3, value=name).border = border

        for col, val in [(4, netto), (5, brutto)]:
            cell = ws.cell(row=row_idx, column=col, value=val)
            cell.border = border
            if val is not None:
                cell.number_format = '#,##0'

        pct_cell = ws.cell(row=row_idx, column=6, value=pct)
        pct_cell.border = border
        if pct is not None:
            pct_cell.number_format = '0.0"%"'

        if is_group and level == 1:
            for c in range(1, 7):
                ws.cell(row=row_idx, column=c).fill = group_fill
                ws.cell(row=row_idx, column=c).font = Font(bold=True)

        row_idx += 1

    # Total row
    total_row = row_idx
    for c, val in enumerate(["", "", "GESAMT", 32340000, 38484700, 100.0], 1):
        cell = ws.cell(row=total_row, column=c, value=val)
        cell.border = border
        cell.fill = total_fill
        cell.font = Font(bold=True, size=12)
        if c in (4, 5) and isinstance(val, (int, float)):
            cell.number_format = '#,##0'
        if c == 6:
            cell.number_format = '0.0"%"'

    ws.column_dimensions["A"].width = 10
    ws.column_dimensions["B"].width = 5
    ws.column_dimensions["C"].width = 38
    ws.column_dimensions["D"].width = 22
    ws.column_dimensions["E"].width = 22
    ws.column_dimensions["F"].width = 14

    wb.save(str(BASE / "04 Kostenplanung" / "Kostenschaetzung_DIN276.xlsx"))


def create_zahlungsplan_xlsx():
    headers = ["Monat", "Geplant (EUR)", "Kumuliert (EUR)", "Ist (EUR)", "Abweichung (EUR)", "Bemerkung"]
    planned = [450000, 980000, 1400000, 1800000, 2200000, 2100000, 1950000, 1800000, 1600000, 1400000]
    months = ["März 2026", "April 2026", "Mai 2026", "Juni 2026", "Juli 2026",
              "August 2026", "September 2026", "Oktober 2026", "November 2026", "Dezember 2026"]
    cum = 0
    rows = []
    for i, (m, p) in enumerate(zip(months, planned)):
        cum += p
        if i == 0:
            ist = 420000
            bemerkung = "Baustelleneinrichtung, erste Erdarbeiten"
        elif i == 1:
            ist = 950000
            bemerkung = "Spundwandarbeiten Nordseite"
        else:
            ist = None
            bemerkung = ""
        abw = (ist - p) if ist else None
        rows.append([m, p, cum, ist, abw, bemerkung])

    make_xlsx(BASE / "04 Kostenplanung" / "Zahlungsplan_2026.xlsx", headers, rows,
              [16, 18, 18, 18, 18, 38])


def create_nachtragsforderungen_docx():
    doc = Document()
    doc.add_heading("Nachtragsforderungen", level=1)
    doc.add_paragraph(
        "Übersicht der eingereichten Nachträge zum Projekt "
        "Neubau TechPark München Ost."
    )

    doc.add_heading("NT-001: Mehrmengen Erdaushub", level=2)
    add_docx_table(doc, ["Feld", "Angabe"], [
        ["Nachtrag-Nr.", "NT-001"],
        ["Auftragnehmer", "Bauer Tiefbau GmbH"],
        ["Gewerk", "Erdarbeiten"],
        ["Datum Eingang", "22.03.2026"],
        ["Forderung netto", "185.000 EUR"],
        ["Status", "In Prüfung"],
    ])
    doc.add_paragraph("")
    doc.add_paragraph(
        "Begründung: Bei den Erdarbeiten im Bereich B wurde eine ca. 1,2 m mächtige "
        "Auffüllungsschicht aus Bauschutt und Betonbruch angetroffen, die im Baugrundgutachten "
        "nicht ausgewiesen war. Die Entsorgung als Z2-Material erfordert einen gesonderten "
        "Entsorgungsweg und verursacht Mehrkosten gegenüber der kalkulierten Klasse Z0."
    )
    p = doc.add_paragraph()
    p.add_run("Prüfvermerk Bauleitung: ").bold = True
    p.add_run(
        "Nachtrag dem Grunde nach berechtigt (geänderter Baugrund). "
        "Prüfung der Höhe steht aus — Mengenaufmaß und Entsorgungsnachweise angefordert."
    )

    doc.add_heading("NT-002: Zusätzliche Verbaumaßnahmen Achse 5–7", level=2)
    add_docx_table(doc, ["Feld", "Angabe"], [
        ["Nachtrag-Nr.", "NT-002"],
        ["Auftragnehmer", "Bauer Tiefbau GmbH"],
        ["Gewerk", "Spezialtiefbau"],
        ["Datum Eingang", "28.03.2026"],
        ["Forderung netto", "312.000 EUR"],
        ["Status", "Eingereicht"],
    ])
    doc.add_paragraph("")
    doc.add_paragraph(
        "Begründung: Aufgrund der höheren Grundwasserstände im Bereich Achse 5–7 "
        "(gemessen: 3,8 m unter GOK statt prognostiziert 4,5 m) sind zusätzliche "
        "Dichtwandarbeiten und eine verstärkte Wasserhaltung erforderlich. "
        "Die ursprünglich geplante offene Wasserhaltung reicht nicht aus."
    )
    p = doc.add_paragraph()
    p.add_run("Prüfvermerk Bauleitung: ").bold = True
    p.add_run(
        "Nachtrag wird geprüft. Grundwassermessungen bestätigen die höheren Stände. "
        "Abstimmung mit Baugrundgutachter und Bauherr erforderlich."
    )

    doc.save(str(BASE / "04 Kostenplanung" / "Nachtragsforderungen.docx"))


# ── 05 PROTOKOLLE ───────────────────────────────────────────────────────

def create_baubesprechung_12_docx():
    doc = Document()
    doc.add_heading("Protokoll Baubesprechung Nr. 12", level=1)

    add_docx_table(doc, ["", ""], [
        ["Datum", "02.04.2026, 10:00 – 12:15 Uhr"],
        ["Ort", "Baucontainer 1, Messestraße 42"],
        ["Protokollführung", "Dipl.-Ing. Thomas Berger"],
        ["Nächster Termin", "09.04.2026, 10:00 Uhr"],
    ])
    doc.add_paragraph("")

    doc.add_heading("Teilnehmer", level=2)
    for name in [
        "Dr. Martin Müller (Bauherr)",
        "Arch. Sabine Lechner (Generalplanerin)",
        "Dipl.-Ing. Thomas Berger (Bauleitung)",
        "Hr. Josef Bauer (Bauer Tiefbau GmbH)",
        "Fr. Claudia Kellner (Züblin AG) — ab TOP 3",
        "Ing. Andrea Hofmann (SiGeKo)",
        "Michael Gruber (Polier)",
    ]:
        doc.add_paragraph(name, style="List Bullet")

    doc.add_heading("TOP 1: Genehmigung Protokoll Nr. 11", level=2)
    doc.add_paragraph(
        "Das Protokoll Nr. 11 vom 26.03.2026 wird ohne Änderungen genehmigt."
    )

    doc.add_heading("TOP 2: Baufortschritt Erdarbeiten", level=2)
    doc.add_paragraph(
        "Hr. Bauer berichtet: Spundwandarbeiten Nordseite (Achse 1–5) zu ca. 80 % abgeschlossen. "
        "Aushub Bereich A planmäßig; Bereich B verzögert durch Altlastenfund (vgl. NT-001). "
        "Voraussichtlicher Abschluss Erdarbeiten: Ende Juni 2026 (unverändert)."
    )
    p = doc.add_paragraph()
    p.add_run("Maßnahme: ").bold = True
    p.add_run("Bauer Tiefbau legt bis 09.04. aktualisierte Terminplanung für Bereich B vor.")

    doc.add_heading("TOP 3: Vorbereitung Rohbau", level=2)
    doc.add_paragraph(
        "Fr. Kellner (Züblin) stellt den Vorentwurf des Schalungskonzepts vor. "
        "Einsatz von Kletterschalung für die Kerne, Deckentische für die Regelgeschosse. "
        "Bewehrungsplanung wird bis KW 18 fertiggestellt."
    )
    p = doc.add_paragraph()
    p.add_run("Maßnahme: ").bold = True
    p.add_run("Lechner-Büro liefert bis 16.04. die Ausführungspläne UG1 (Maßstab 1:50).")

    doc.add_heading("TOP 4: Nachträge", level=2)
    doc.add_paragraph(
        "NT-001 (Mehrmengen Erdaushub): Mengenaufmaß liegt vor, Prüfung durch Bauleitung "
        "bis 09.04. NT-002 (Zusätzlicher Verbau): Besprechung mit Baugrundgutachter am 04.04."
    )

    doc.add_heading("TOP 5: Sicherheit und Verschiedenes", level=2)
    doc.add_paragraph(
        "Fr. Hofmann (SiGeKo): Begehung am 01.04. ohne wesentliche Beanstandungen. "
        "Hinweis auf fehlende Absturzsicherung am Grubenrand Bereich C — bis 03.04. nachzurüsten."
    )
    p = doc.add_paragraph()
    p.add_run("Maßnahme: ").bold = True
    p.add_run("Gruber stellt Absturzsicherung Bereich C bis 03.04. sicher.")

    doc.save(str(BASE / "05 Protokolle" / "Baubesprechung_Nr12_2026-04-02.docx"))


def create_baubesprechung_11_docx():
    doc = Document()
    doc.add_heading("Protokoll Baubesprechung Nr. 11", level=1)

    add_docx_table(doc, ["", ""], [
        ["Datum", "26.03.2026, 10:00 – 11:30 Uhr"],
        ["Ort", "Baucontainer 1, Messestraße 42"],
        ["Protokollführung", "Dipl.-Ing. Thomas Berger"],
        ["Nächster Termin", "02.04.2026, 10:00 Uhr"],
    ])
    doc.add_paragraph("")

    doc.add_heading("Teilnehmer", level=2)
    for name in [
        "Dr. Martin Müller (Bauherr)",
        "Arch. Sabine Lechner (Generalplanerin)",
        "Dipl.-Ing. Thomas Berger (Bauleitung)",
        "Hr. Josef Bauer (Bauer Tiefbau GmbH)",
        "Ing. Andrea Hofmann (SiGeKo)",
        "Michael Gruber (Polier)",
    ]:
        doc.add_paragraph(name, style="List Bullet")

    doc.add_heading("TOP 1: Baufortschritt", level=2)
    doc.add_paragraph(
        "Spundwandarbeiten Achse 1–3 abgeschlossen. Achse 3–5 in Arbeit. "
        "Erdaushub Bereich A läuft planmäßig. Altlastenfund in Bereich B — "
        "Entsorgungskonzept wird erarbeitet."
    )

    doc.add_heading("TOP 2: Nachtrag NT-001", level=2)
    doc.add_paragraph(
        "Hr. Bauer hat Nachtragsforderung NT-001 eingereicht (185.000 EUR netto). "
        "Bauleitung prüft den Nachtrag. Mengenaufmaß wird bis 02.04. nachgereicht."
    )

    doc.add_heading("TOP 3: Sicherheit", level=2)
    doc.add_paragraph(
        "SiGeKo-Begehung am 25.03.: Bauzaunfeld Nr. 14 beschädigt — wurde bereits repariert. "
        "Keine weiteren Mängel."
    )

    doc.save(str(BASE / "05 Protokolle" / "Baubesprechung_Nr11_2026-03-26.docx"))


def create_abnahmeprotokoll_docx():
    doc = Document()
    doc.add_heading("Abnahmeprotokoll — Vorlage", level=1)
    doc.add_paragraph(
        "Projekt: Neubau Bürogebäude TechPark München Ost"
    )

    add_docx_table(doc, ["Feld", "Eintrag"], [
        ["Gewerk / Leistung", "________________________________________"],
        ["Auftragnehmer", "________________________________________"],
        ["Datum der Abnahme", "________________________________________"],
        ["Ort / Bereich", "________________________________________"],
        ["Teilnehmer AN", "________________________________________"],
        ["Teilnehmer AG / BL", "________________________________________"],
    ])
    doc.add_paragraph("")

    doc.add_heading("Ergebnis der Abnahme", level=2)
    doc.add_paragraph("☐  Abnahme erklärt (ohne Mängel)")
    doc.add_paragraph("☐  Abnahme erklärt (mit Mängeln lt. Mängelliste)")
    doc.add_paragraph("☐  Abnahme verweigert (wesentliche Mängel)")

    doc.add_heading("Festgestellte Mängel", level=2)
    add_docx_table(doc, ["Nr.", "Beschreibung", "Frist", "Bemerkung"], [
        ["1", "", "", ""],
        ["2", "", "", ""],
        ["3", "", "", ""],
        ["4", "", "", ""],
        ["5", "", "", ""],
    ])

    doc.add_heading("Vereinbarungen", level=2)
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")

    doc.add_heading("Unterschriften", level=2)
    doc.add_paragraph("")
    doc.add_paragraph("_________________________          _________________________")
    doc.add_paragraph("Auftragnehmer                              Auftraggeber / Bauleitung")
    doc.add_paragraph("")
    doc.add_paragraph("Ort, Datum: _________________________")

    doc.save(str(BASE / "05 Protokolle" / "Abnahmeprotokoll_Vorlage.docx"))


# ── 06 FOTOS (SVG) ──────────────────────────────────────────────────────

def create_svg_placeholder(filename, bg_color, text, sub_text=""):
    svg = f'''<?xml version="1.0" encoding="UTF-8"?>
<svg xmlns="http://www.w3.org/2000/svg" width="800" height="600" viewBox="0 0 800 600">
  <rect width="800" height="600" fill="{bg_color}" rx="8"/>
  <rect x="50" y="50" width="700" height="500" fill="rgba(255,255,255,0.15)" rx="4" stroke="rgba(255,255,255,0.3)" stroke-width="2"/>
  <!-- Camera icon -->
  <g transform="translate(350, 200)" fill="rgba(255,255,255,0.4)">
    <rect x="-40" y="-25" width="80" height="55" rx="8"/>
    <circle cx="0" cy="5" r="18" fill="none" stroke="rgba(255,255,255,0.4)" stroke-width="3"/>
    <rect x="-12" y="-32" width="24" height="10" rx="3"/>
  </g>
  <text x="400" y="340" font-family="Arial, Helvetica, sans-serif" font-size="28" font-weight="bold"
        fill="white" text-anchor="middle">{text}</text>
  <text x="400" y="375" font-family="Arial, Helvetica, sans-serif" font-size="16"
        fill="rgba(255,255,255,0.7)" text-anchor="middle">{sub_text}</text>
  <text x="400" y="520" font-family="Arial, Helvetica, sans-serif" font-size="14"
        fill="rgba(255,255,255,0.5)" text-anchor="middle">Platzhalter — Baudokumentation TechPark München Ost</text>
</svg>'''
    (BASE / "06 Fotos" / filename).write_text(svg, encoding="utf-8")


def create_all_svgs():
    create_svg_placeholder(
        "Baustellenuebersicht_2026-03-10.svg", "#2563EB",
        "Baustellenübersicht", "10.03.2026 — Gesamtansicht von Südost"
    )
    create_svg_placeholder(
        "Spundwand_Achse1-3_2026-03-25.svg", "#16A34A",
        "Spundwandarbeiten Achse 1–3", "25.03.2026 — Nordseite"
    )
    create_svg_placeholder(
        "Baugrube_von_oben_2026-04-01.svg", "#7C3AED",
        "Baugrube — Drohnenaufnahme", "01.04.2026 — Luftbild"
    )
    create_svg_placeholder(
        "Baustellenzufahrt_2026-03-05.svg", "#92400E",
        "Baustelleneinrichtung / Zufahrt Tor 1", "05.03.2026 — Messestraße"
    )


# ── 07 VERTRÄGE ──────────────────────────────────────────────────────────

def create_generalplanervertrag_docx():
    doc = Document()
    doc.add_heading("Generalplanervertrag — Entwurf", level=1)
    doc.add_paragraph(
        'zwischen M\u00fcller & Schmidt Immobilien GmbH (nachfolgend \u201eAuftraggeber\u201c) '
        'und Architekten Lechner & Kollegen PartGmbB (nachfolgend \u201eAuftragnehmer\u201c)'
    )

    doc.add_heading("Präambel", level=2)
    doc.add_paragraph(
        "Der Auftraggeber beabsichtigt die Errichtung eines Bürogebäudes "
        "auf dem Grundstück Messestraße 42, 81829 München. Der Auftragnehmer "
        "wird als Generalplaner mit der umfassenden Planungsleistung beauftragt."
    )

    doc.add_heading("§ 1 Leistungsumfang", level=2)
    doc.add_paragraph(
        "Der Auftragnehmer übernimmt folgende Leistungen gemäß HOAI 2021:"
    )
    for item in [
        "Objektplanung Gebäude (LPH 1–9)",
        "Tragwerksplanung (LPH 1–6), ggf. als Subplaner",
        "Technische Ausrüstung — Koordination der Fachplaner (LPH 1–9)",
        "Freianlagenplanung (LPH 1–9)",
        "Brandschutzplanung (als Subplaner)",
        "Koordination aller Planungsbeteiligten",
        "Mitwirkung bei der Vergabe (LPH 6–7)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("§ 2 Honorar", level=2)
    doc.add_paragraph(
        "Das Honorar richtet sich nach der HOAI 2021 und wird auf Basis der "
        "anrechenbaren Kosten wie folgt vereinbart:"
    )
    add_docx_table(doc, ["Position", "Betrag (EUR netto)"], [
        ["Objektplanung Gebäude (LPH 1–9)", "2.450.000"],
        ["Tragwerksplanung (LPH 1–6)", "380.000"],
        ["TGA-Koordination", "210.000"],
        ["Freianlagenplanung", "95.000"],
        ["Brandschutzplanung", "65.000"],
        ["Besondere Leistungen (Pauschale)", "120.000"],
        ["Gesamt netto", "3.320.000"],
    ])

    doc.add_heading("§ 3 Termine", level=2)
    add_docx_table(doc, ["Leistungsphase", "Fertigstellung bis"], [
        ["LPH 1–2 (Grundlagenermittlung, Vorplanung)", "Abgeschlossen"],
        ["LPH 3 (Entwurfsplanung)", "Abgeschlossen"],
        ["LPH 4 (Genehmigungsplanung)", "Abgeschlossen (Genehmigung erteilt 15.02.2026)"],
        ["LPH 5 (Ausführungsplanung)", "Laufend — Fertigstellung bis 30.06.2026"],
        ["LPH 6–7 (Vergabe)", "Laufend — Hauptvergaben bis 31.03.2026"],
        ["LPH 8 (Bauüberwachung)", "Baubegleitend bis Abnahme"],
        ["LPH 9 (Objektbetreuung)", "5 Jahre nach Abnahme"],
    ])

    doc.save(str(BASE / "07 Vertraege" / "Generalplanervertrag_Entwurf.docx"))


def create_buergschaftsuebersicht_xlsx():
    headers = ["Firma", "Art der Bürgschaft", "Betrag (EUR)", "Bank", "Gültig bis", "Status"]
    rows = [
        ["Bauer Tiefbau GmbH", "Vertragserfüllungsbürgschaft", 285000, "Sparkasse München", "31.12.2026", "Vorliegend"],
        ["Züblin AG, NL München", "Vertragserfüllungsbürgschaft", 820000, "Deutsche Bank AG", "31.12.2027", "Vorliegend"],
        ["Züblin AG, NL München", "Gewährleistungsbürgschaft", 410000, "Deutsche Bank AG", "30.09.2033", "Noch nicht fällig"],
        ["Haustechnik Süd GmbH", "Vertragserfüllungsbürgschaft", 380000, "Volksbank München", "31.12.2028", "Angefordert"],
        ["Elektro Müller GmbH", "Vertragserfüllungsbürgschaft", 210000, "Commerzbank AG", "31.12.2028", "Angefordert"],
    ]
    make_xlsx(BASE / "07 Vertraege" / "Buergschaftsuebersicht.xlsx", headers, rows,
              [28, 30, 18, 22, 16, 18])


# ── 08 ARBEITSSICHERHEIT ────────────────────────────────────────────────

def create_sigeplan_docx():
    doc = Document()
    doc.add_heading("SiGe-Plan — Übersicht", level=1)
    doc.add_paragraph(
        "Sicherheits- und Gesundheitsschutzplan gemäß BaustellV "
        "für das Vorhaben Neubau TechPark München Ost."
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("SiGeKo: ").bold = True
    p.add_run("Ing. Andrea Hofmann, Hofmann Sicherheitstechnik")
    p = doc.add_paragraph()
    p.add_run("Erstellungsdatum: ").bold = True
    p.add_run("28.02.2026")
    p = doc.add_paragraph()
    p.add_run("Fortschreibung: ").bold = True
    p.add_run("Laufend, jeweils bei Bauphasenanpassung")

    doc.add_heading("Gefährdungsbeurteilung — Bauphase Erdarbeiten", level=2)
    add_docx_table(doc, ["Nr.", "Gefährdung", "Maßnahme", "Verantwortlich", "Termin"], [
        ["G-01", "Absturz in Baugrube (Tiefe > 5 m)", "Absturzsicherung gem. TRBS 2121; Seitenschutz mind. 1,0 m", "Bauer Tiefbau / Polier", "Vor Aushub"],
        ["G-02", "Verschüttung durch Böschungsbruch", "Böschungswinkel einhalten (max. 60°); ggf. Verbau", "Bauer Tiefbau", "Laufend"],
        ["G-03", "Gefährdung durch Baumaschinenverkehr", "Einweiser bei Rückwärtsfahrten; Sicherheitsabstände", "Polier", "Laufend"],
        ["G-04", "Lärm (Rammarbeiten Spundwand)", "Gehörschutz ab 85 dB(A); Lärmbereich kennzeichnen", "SiGeKo / Polier", "Während Rammarbeiten"],
        ["G-05", "Kontamination durch Altlasten", "PSA (Handschuhe, Atemschutz bei Bedarf); Bodengutachter hinzuziehen", "Bauer Tiefbau", "Bei Altlastenfund"],
    ])

    doc.add_heading("Persönliche Schutzausrüstung (PSA)", level=2)
    doc.add_paragraph("Grundausstattung für alle Personen auf der Baustelle:")
    for item in [
        "Schutzhelm (EN 397) — Pflicht auf dem gesamten Baufeld",
        "Sicherheitsschuhe S3 (EN ISO 20345)",
        "Warnweste Klasse 2 (EN ISO 20471)",
        "Schutzhandschuhe bei manuellen Arbeiten",
        "Gehörschutz bei Lärmarbeiten (> 85 dB(A))",
        "Schutzbrille bei Schneid-, Schweiß- und Schleifarbeiten",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Notfallplan", level=2)
    add_docx_table(doc, ["", ""], [
        ["Sammelplatz", "Parkplatz Nordost (gekennzeichnet mit grünem Schild)"],
        ["Ersthelfer", "Michael Gruber (Polier) — 0170 333 9876"],
        ["Nächstes Krankenhaus", "Klinikum Bogenhausen — Englschalkinger Str. 77 (ca. 10 min)"],
        ["Notruf", "112"],
        ["Feuerwehr (direkt)", "089 / 2353-0"],
        ["Giftnotruf München", "089 / 19240"],
    ])

    doc.save(str(BASE / "08 Arbeitssicherheit" / "SiGePlan_Uebersicht.docx"))


def create_unterweisungsnachweis_xlsx():
    headers = ["Datum", "Name", "Firma", "Gewerk", "Thema", "Unterweisung durch", "Unterschrift"]
    rows = [
        ["01.03.2026", "Josef Bauer", "Bauer Tiefbau GmbH", "Erdarbeiten", "Baustellenordnung / Allgemeine Sicherheit", "Ing. Andrea Hofmann", "✓"],
        ["01.03.2026", "Michael Gruber", "Bauer Tiefbau GmbH", "Polier", "Baustellenordnung / Allgemeine Sicherheit", "Ing. Andrea Hofmann", "✓"],
        ["01.03.2026", "Markus Steiner", "Bauer Tiefbau GmbH", "Erdarbeiten", "Baustellenordnung / Allgemeine Sicherheit", "Ing. Andrea Hofmann", "✓"],
        ["01.03.2026", "Andreas Lang", "Bauer Tiefbau GmbH", "Erdarbeiten", "Baustellenordnung / Allgemeine Sicherheit", "Ing. Andrea Hofmann", "✓"],
        ["05.03.2026", "Josef Bauer", "Bauer Tiefbau GmbH", "Erdarbeiten", "Arbeiten in Baugruben / Absturzgefahr", "Ing. Andrea Hofmann", "✓"],
        ["05.03.2026", "Michael Gruber", "Bauer Tiefbau GmbH", "Polier", "Arbeiten in Baugruben / Absturzgefahr", "Ing. Andrea Hofmann", "✓"],
        ["10.03.2026", "Peter Krämer", "Vermessungsbüro Krämer", "Vermessung", "Baustellenordnung / Allgemeine Sicherheit", "Michael Gruber", "✓"],
        ["15.03.2026", "Josef Bauer", "Bauer Tiefbau GmbH", "Spezialtiefbau", "Spundwandarbeiten / Lärmschutz / PSA", "Ing. Andrea Hofmann", "✓"],
        ["15.03.2026", "Markus Steiner", "Bauer Tiefbau GmbH", "Spezialtiefbau", "Spundwandarbeiten / Lärmschutz / PSA", "Ing. Andrea Hofmann", "✓"],
        ["20.03.2026", "Thomas Berger", "Berger Baumanagement", "Bauleitung", "Umgang mit kontaminierten Böden (Altlastenfund)", "Ing. Andrea Hofmann", "✓"],
    ]
    make_xlsx(BASE / "08 Arbeitssicherheit" / "Unterweisungsnachweis_2026-03.xlsx", headers, rows,
              [14, 20, 26, 18, 42, 22, 14])


def create_unfallbericht_docx():
    doc = Document()
    doc.add_heading("Unfallbericht — Vorlage", level=1)
    doc.add_paragraph("Projekt: Neubau Bürogebäude TechPark München Ost")
    doc.add_paragraph("")

    doc.add_heading("Angaben zum Unfall", level=2)
    add_docx_table(doc, ["Feld", "Eintrag"], [
        ["Datum des Unfalls", "________________________________________"],
        ["Uhrzeit", "________________________________________"],
        ["Ort (genau)", "________________________________________"],
        ["Name des Verletzten", "________________________________________"],
        ["Firma", "________________________________________"],
        ["Gewerk / Tätigkeit", "________________________________________"],
    ])

    doc.add_heading("Unfallhergang", level=2)
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")

    doc.add_heading("Art der Verletzung", level=2)
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")

    doc.add_heading("Erste Hilfe / Maßnahmen", level=2)
    doc.add_paragraph("☐  Erstversorgung vor Ort durch: ________________________")
    doc.add_paragraph("☐  Rettungsdienst gerufen (112)")
    doc.add_paragraph("☐  Transport ins Krankenhaus: ________________________")
    doc.add_paragraph("☐  D-Arzt-Besuch veranlasst")

    doc.add_heading("Zeugen", level=2)
    add_docx_table(doc, ["Name", "Firma", "Telefon"], [
        ["", "", ""],
        ["", "", ""],
    ])

    doc.add_heading("Sofortmaßnahmen / Ursachenanalyse", level=2)
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")
    doc.add_paragraph("____________________________________________________________")

    doc.add_heading("Meldung", level=2)
    doc.add_paragraph("☐  BG BAU informiert (bei Arbeitsunfähigkeit > 3 Tage)")
    doc.add_paragraph("☐  Bauherr / Bauleitung informiert")
    doc.add_paragraph("☐  SiGeKo informiert")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("_________________________          _________________________")
    doc.add_paragraph("Erstellt von (Name, Datum)           SiGeKo (Unterschrift)")

    doc.save(str(BASE / "08 Arbeitssicherheit" / "Unfallbericht_Vorlage.docx"))


# ── MAIN ─────────────────────────────────────────────────────────────────

def main():
    print("Creating directory structure...")
    ensure_dirs()

    generators = [
        ("Projektuebersicht.md", create_projektuebersicht_md),
        ("Kontaktliste.xlsx", create_kontaktliste_xlsx),
        ("01 Planung/Projektbeschreibung.docx", create_projektbeschreibung_docx),
        ("01 Planung/Raumprogramm.xlsx", create_raumprogramm_xlsx),
        ("01 Planung/Terminplan_Uebersicht.pdf", create_terminplan_pdf),
        ("01 Planung/Architekturbriefing.docx", create_architekturbriefing_docx),
        ("02 Genehmigungen/Baugenehmigung_Bescheid.pdf", create_baugenehmigung_pdf),
        ("02 Genehmigungen/Umweltvertraeglichkeitspruefung.docx", create_umweltvertraeglichkeit_docx),
        ("02 Genehmigungen/Brandschutzkonzept_Entwurf.docx", create_brandschutzkonzept_docx),
        ("03 Bauausfuehrung/Baustellenordnung.pdf", create_baustellenordnung_pdf),
        ("03 Bauausfuehrung/Nachunternehmer_Uebersicht.xlsx", create_nachunternehmer_xlsx),
        ("03 Bauausfuehrung/Bautagesbericht_2026-03-15.docx", create_bautagesbericht_docx),
        ("03 Bauausfuehrung/Maengelliste.xlsx", create_maengelliste_xlsx),
        ("04 Kostenplanung/Kostenschaetzung_DIN276.xlsx", create_kostenschaetzung_xlsx),
        ("04 Kostenplanung/Zahlungsplan_2026.xlsx", create_zahlungsplan_xlsx),
        ("04 Kostenplanung/Nachtragsforderungen.docx", create_nachtragsforderungen_docx),
        ("05 Protokolle/Baubesprechung_Nr12_2026-04-02.docx", create_baubesprechung_12_docx),
        ("05 Protokolle/Baubesprechung_Nr11_2026-03-26.docx", create_baubesprechung_11_docx),
        ("05 Protokolle/Abnahmeprotokoll_Vorlage.docx", create_abnahmeprotokoll_docx),
        ("06 Fotos/*.svg", create_all_svgs),
        ("07 Vertraege/Generalplanervertrag_Entwurf.docx", create_generalplanervertrag_docx),
        ("07 Vertraege/Buergschaftsuebersicht.xlsx", create_buergschaftsuebersicht_xlsx),
        ("08 Arbeitssicherheit/SiGePlan_Uebersicht.docx", create_sigeplan_docx),
        ("08 Arbeitssicherheit/Unterweisungsnachweis_2026-03.xlsx", create_unterweisungsnachweis_xlsx),
        ("08 Arbeitssicherheit/Unfallbericht_Vorlage.docx", create_unfallbericht_docx),
    ]

    for name, func in generators:
        try:
            func()
            print(f"  OK: {name}")
        except Exception as e:
            print(f"  FAIL: {name}: {e}")

    # Verify
    print("\n--- Verification ---")
    total = 0
    for root, dirs, files in os.walk(BASE):
        for f in files:
            total += 1
            rel = os.path.relpath(os.path.join(root, f), BASE)
            size = os.path.getsize(os.path.join(root, f))
            print(f"  {rel} ({size:,} bytes)")
    print(f"\nTotal files: {total}")


if __name__ == "__main__":
    main()
