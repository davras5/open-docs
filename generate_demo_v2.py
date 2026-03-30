"""
Generate professional demo files for OpenDocs — Swiss construction project.
Überbauung Seefeld, Zürich (mixed-use residential/commercial).
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from fpdf import FPDF

BASE = os.path.dirname(os.path.abspath(__file__))
DEMO = os.path.join(BASE, "demo-files")

# ── Helpers ──

def ensure_dir(path):
    os.makedirs(os.path.join(DEMO, path), exist_ok=True)

def styled_xlsx(path, headers, rows, col_widths=None, sheet_name="Sheet1"):
    """Create a clean, professional Excel file."""
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    hfont = Font(bold=True, size=10, name="Calibri")
    nfont = Font(size=10, name="Calibri")
    hfill = PatternFill(start_color="F2F4F7", end_color="F2F4F7", fill_type="solid")
    thin = Side(style="thin", color="D0D5DD")
    border = Border(top=thin, bottom=thin, left=thin, right=thin)

    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = hfont
        cell.fill = hfill
        cell.border = border
        cell.alignment = Alignment(horizontal="left", vertical="center")

    for r, row in enumerate(rows, 2):
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = nfont
            cell.border = border
            cell.alignment = Alignment(vertical="center")

    if col_widths:
        for i, w in enumerate(col_widths):
            ws.column_dimensions[chr(65 + i)].width = w
    else:
        for i in range(len(headers)):
            ws.column_dimensions[chr(65 + i)].width = max(15, len(headers[i]) + 8)

    ws.auto_filter.ref = ws.dimensions
    ws.freeze_panes = "A2"
    wb.save(os.path.join(DEMO, path))
    print(f"  ✓ {path}")


# ══════════════════════════════════════════════════════════════════════════════
# Clean old files
# ══════════════════════════════════════════════════════════════════════════════
import shutil
for item in os.listdir(DEMO):
    p = os.path.join(DEMO, item)
    if os.path.isdir(p):
        shutil.rmtree(p)
    elif item != ".gitkeep":
        os.remove(p)

# Create folder structure
folders = [
    "01 Planung",
    "02 Bewilligungen",
    "03 Bauausführung",
    "04 Kosten",
    "05 Protokolle",
    "06 Fotos",
    "07 Verträge",
    "08 SIBE",
]
for f in folders:
    ensure_dir(f)

print("Generating Swiss construction project demo files...\n")

# ══════════════════════════════════════════════════════════════════════════════
# ROOT FILES
# ══════════════════════════════════════════════════════════════════════════════

# Projektübersicht.md
with open(os.path.join(DEMO, "Projektübersicht.md"), "w", encoding="utf-8") as f:
    f.write("""# Überbauung Seefeld — Neubau Wohn- und Geschäftshaus Zürich

## Projektdaten

| Eigenschaft | Wert |
|---|---|
| Bauherrschaft | Seefeld Immobilien AG, Zürich |
| Standort | Seefeldstrasse 128, 8008 Zürich |
| Generalplaner | Baumschlager Eberle Architekten, Zürich |
| Bauleitung | Implenia Schweiz AG, Dietikon |
| Baubeginn | Februar 2026 |
| Bezug | Juni 2028 (geplant) |
| Anlagekosten BKP 1–9 | CHF 62,8 Mio. |

## Projektbeschrieb

Die Überbauung Seefeld umfasst den Neubau eines siebengeschossigen Wohn- und
Geschäftshauses mit Tiefgarage an der Seefeldstrasse 128 in Zürich. Das Gebäude
bietet 48 Mietwohnungen (2.5–5.5 Zimmer), Gewerbeflächen im Erdgeschoss sowie
68 Tiefgaragenplätze. Der Entwurf setzt auf eine hochwertige Fassade aus
Sichtbeton und Holz-Metall-Fenstern, begrünte Innenhöfe und Dachterrassen.

## Ordnerstruktur

- **01 Planung** — Projektbeschrieb, Raumprogramm, Terminplan
- **02 Bewilligungen** — Baubewilligung, UVB, Brandschutz
- **03 Bauausführung** — Baustellenordnung, Unternehmerverzeichnis
- **04 Kosten** — BKP-Kostenvoranschlag, Zahlungsplan, Nachträge
- **05 Protokolle** — Bausitzungen, Abnahmen
- **06 Fotos** — Baudokumentation
- **07 Verträge** — Planervertrag, Werkverträge
- **08 SIBE** — Sicherheitskonzept, Unterweisungen

## Projektbeteiligte

| Name | Funktion | Kontakt |
|---|---|---|
| Dr. Stefan Meier | Bauherrenvertreter | s.meier@seefeld-immo.ch |
| Arch. ETH Laura Brunner | Projektleiterin GP | brunner@be-arch.ch |
| Dipl. Bauing. ETH Marco Keller | Bauleiter | m.keller@implenia.com |
| Ing. Nina Frei | Sicherheitsbeauftragte | n.frei@sibe-zuerich.ch |
| Roman Steiner | Polier | r.steiner@implenia.com |
""")
print("  ✓ Projektübersicht.md")

# Kontaktliste.xlsx
styled_xlsx("Kontaktliste.xlsx",
    ["Name", "Firma", "Funktion", "E-Mail", "Telefon"],
    [
        ["Dr. Stefan Meier", "Seefeld Immobilien AG", "Bauherrenvertreter", "s.meier@seefeld-immo.ch", "+41 44 380 22 10"],
        ["Arch. ETH Laura Brunner", "Baumschlager Eberle Architekten", "Projektleiterin GP", "brunner@be-arch.ch", "+41 44 561 33 00"],
        ["Dipl. Bauing. ETH Marco Keller", "Implenia Schweiz AG", "Bauleiter", "m.keller@implenia.com", "+41 58 474 00 00"],
        ["Ing. Nina Frei", "SIBE Zürich GmbH", "Sicherheitsbeauftragte", "n.frei@sibe-zuerich.ch", "+41 44 720 15 30"],
        ["Roman Steiner", "Implenia Schweiz AG", "Polier", "r.steiner@implenia.com", "+41 79 445 12 88"],
        ["Dr. Ing. Thomas Wyss", "Schnetzer Puskas Ingenieure", "Tragwerksplaner", "wyss@schnetzer-puskas.ch", "+41 44 315 77 00"],
        ["Markus Roth", "Amstein + Walthert AG", "HLKS-Planer", "roth@amstein-walthert.ch", "+41 44 305 91 11"],
        ["Claudia Gerber", "Kanton Zürich, Baupolizei", "Baubehörde", "claudia.gerber@bd.zh.ch", "+41 43 259 30 45"],
        ["Peter Huber", "Marti AG Zürich", "Baumeisterarbeiten", "p.huber@martiag.ch", "+41 44 307 80 00"],
        ["Sandra Bieri", "Ernst Basler + Partner", "Umweltberatung", "bieri@ebp.ch", "+41 44 395 11 11"],
    ],
    col_widths=[30, 35, 22, 30, 20]
)

# ══════════════════════════════════════════════════════════════════════════════
# 01 PLANUNG
# ══════════════════════════════════════════════════════════════════════════════

# Projektbeschrieb.docx (multi-page)
doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Calibri'
style.paragraph_format.space_after = Pt(6)

doc.add_heading('Projektbeschrieb', level=0)
doc.add_paragraph('Überbauung Seefeld — Neubau Wohn- und Geschäftshaus')
doc.add_paragraph('Seefeldstrasse 128, 8008 Zürich')
doc.add_paragraph('')

doc.add_heading('1. Ausgangslage', level=1)
doc.add_paragraph(
    'Auf dem Grundstück Kat.-Nr. 8008-1247 an der Seefeldstrasse 128 in Zürich befindet sich '
    'derzeit ein zweigeschossiges Gewerbegebäude aus den 1960er-Jahren, das nicht mehr den '
    'heutigen Anforderungen an Energieeffizienz und Nutzung entspricht. Das Grundstück liegt '
    'in der Wohnzone W5 gemäss BZO der Stadt Zürich und ist dem ISOS-Gebiet «Seefeld» '
    'zugeordnet. Die Seefeld Immobilien AG beabsichtigt den Abbruch des Bestandsgebäudes '
    'und die Errichtung eines zeitgemässen Wohn- und Geschäftshauses.'
)

doc.add_heading('2. Projektkonzept', level=1)
doc.add_paragraph(
    'Der Neubau umfasst sieben Obergeschosse, ein Attikageschoss mit Dachterrassen '
    'sowie zwei Untergeschosse für Tiefgarage und Haustechnik. Die Architektur orientiert '
    'sich an der für das Seefeld-Quartier typischen Blockrandbebauung und integriert '
    'einen begrünten Innenhof, der allen Bewohnern als gemeinschaftlicher Freiraum dient.'
)

doc.add_heading('2.1 Nutzungskonzept', level=2)
table = doc.add_table(rows=9, cols=4)
table.style = 'Table Grid'
for i, h in enumerate(['Geschoss', 'Nutzung', 'Fläche (m²)', 'Bemerkung']):
    table.rows[0].cells[i].text = h
data = [
    ['EG', 'Gewerbe / Detailhandel', '680', 'Drei Ladenlokale, Zugang Innenhof'],
    ['1.–3. OG', 'Mietwohnungen', '3 × 920', '18 Wohnungen (2.5–4.5 Zi.)'],
    ['4.–5. OG', 'Mietwohnungen', '2 × 880', '12 Wohnungen (3.5–5.5 Zi.)'],
    ['6. OG', 'Mietwohnungen', '760', '6 Wohnungen (4.5–5.5 Zi.)'],
    ['Attika', 'Mietwohnungen', '520', '4 Attika-Wohnungen mit Terrasse'],
    ['1. UG', 'Tiefgarage', '2 400', '68 Parkplätze, Veloraum'],
    ['2. UG', 'Haustechnik / Keller', '1 800', 'Technikzentrale, Kellerabteile'],
    ['Dach', 'PV-Anlage / Begrünung', '950', 'Biodiversitätsdach'],
]
for r, row in enumerate(data, 1):
    for c, val in enumerate(row):
        table.rows[r].cells[c].text = val

doc.add_heading('2.2 Konstruktion und Materialisierung', level=2)
doc.add_paragraph(
    'Tragstruktur: Stahlbetonskelettkonstruktion mit aussteifenden Kernen (Treppenhäuser, '
    'Liftschächte). Geschossdecken als Flachdecken d=26 cm. Gründung auf Bohrpfählen '
    '(Ø 80 cm, Länge ca. 18 m bis Obere Süsswassermolasse).'
)
doc.add_paragraph(
    'Fassade: Vorfabrizierte Sichtbetonelemente (Besenstrich-Struktur, Farbton «Sandstein hell») '
    'mit Holz-Metall-Fenstern (3-fach-Verglasung, Uw ≤ 0.9 W/m²K). Aussenliegender, '
    'textiler Sonnenschutz mit Windautomatik.'
)
doc.add_paragraph(
    'Innenausbau: Anhydrit-Unterlagsboden mit Fussbodenheizung, Eichenparkett in Wohnräumen, '
    'Feinsteinzeug in Nasszellen. Küchen mit Naturstein-Abdeckung und Markengeräten.'
)

doc.add_heading('2.3 Energie und Nachhaltigkeit', level=2)
doc.add_paragraph(
    'Das Gebäude wird nach dem Standard Minergie-P-ECO zertifiziert. Die Wärmeversorgung '
    'erfolgt über Erdsonden-Wärmepumpen (16 Bohrungen à 250 m). Auf der Dachfläche wird '
    'eine Photovoltaik-Anlage mit 180 kWp installiert, die zusammen mit einer Batterie '
    '(150 kWh) den Eigenverbrauch optimiert. Eine kontrollierte Wohnungslüftung mit '
    'Wärmerückgewinnung (η ≥ 82%) ist in allen Wohneinheiten vorgesehen.'
)
items = [
    'Heizwärmebedarf: ≤ 20 kWh/m²a',
    'Primärenergie nicht erneuerbar: ≤ 60 kWh/m²a',
    'CO₂-Emissionen: ≤ 4.0 kg/m²a',
    'Eigenversorgungsgrad Strom: ca. 45%',
    'Graue Energie: Optimierung durch Recyclingbeton (RC-C 30/37)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3. Umgebungsgestaltung', level=1)
doc.add_paragraph(
    'Der Innenhof (ca. 320 m²) wird als naturnaher Begegnungsraum gestaltet mit heimischen '
    'Baumarten (Feldahorn, Hainbuche), extensiver Begrünung und einem Spielbereich. Die '
    'Zufahrt zur Tiefgarage erfolgt über die Seitenstrasse (Bellerivestrasse). Die Vorzonen '
    'entlang der Seefeldstrasse werden als Sitzplätze für die Gewerbeeinheiten genutzt.'
)

doc.add_heading('4. Termine und Budget', level=1)
doc.add_paragraph(
    'Die Bauzeit beträgt voraussichtlich 28 Monate. Der Bezug der Wohnungen ist für '
    'Juni 2028 geplant. Die Anlagekosten (BKP 1–9) belaufen sich auf CHF 62,8 Mio.'
)
table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Table Grid'
for i, h in enumerate(['Phase', 'Zeitraum', 'Dauer']):
    table2.rows[0].cells[i].text = h
phases = [
    ['Abbruch Bestandsgebäude', 'Feb – Mär 2026', '6 Wochen'],
    ['Baugrube und Spezialtiefbau', 'Mär – Jun 2026', '14 Wochen'],
    ['Rohbau', 'Jul 2026 – Apr 2027', '10 Monate'],
    ['Fassade und Dach', 'Mai – Sep 2027', '5 Monate'],
    ['Innenausbau und Haustechnik', 'Okt 2027 – Apr 2028', '7 Monate'],
    ['Umgebung und Abnahmen', 'Mai – Jun 2028', '8 Wochen'],
]
for r, row in enumerate(phases, 1):
    for c, val in enumerate(row):
        table2.rows[r].cells[c].text = val

doc.save(os.path.join(DEMO, "01 Planung", "Projektbeschrieb.docx"))
print("  ✓ 01 Planung/Projektbeschrieb.docx")

# Raumprogramm.xlsx
styled_xlsx("01 Planung/Raumprogramm.xlsx",
    ["Geschoss", "Bereich", "Fläche (m²)", "Wohnungen", "Zimmer", "Bemerkung"],
    [
        ["EG", "Gewerbe A", 280, None, None, "Ladenlokal Seefeldstrasse"],
        ["EG", "Gewerbe B", 220, None, None, "Ladenlokal Seefeldstrasse"],
        ["EG", "Gewerbe C", 180, None, None, "Gastro / Café mit Terrasse"],
        ["EG", "Eingangsbereich", 85, None, None, "Lobby, Briefkästen, Veloabstell"],
        ["1. OG", "Wohnungen", 920, 6, "2.5–4.5", ""],
        ["2. OG", "Wohnungen", 920, 6, "2.5–4.5", ""],
        ["3. OG", "Wohnungen", 920, 6, "2.5–4.5", ""],
        ["4. OG", "Wohnungen", 880, 6, "3.5–5.5", "Grösserer Grundriss"],
        ["5. OG", "Wohnungen", 880, 6, "3.5–5.5", ""],
        ["6. OG", "Wohnungen", 760, 6, "4.5–5.5", ""],
        ["Attika", "Wohnungen", 520, 4, "4.5–5.5", "Mit Dachterrassen"],
        ["Attika", "Gemeinschaftsterrasse", 180, None, None, "Zugänglich für alle Bewohner"],
        ["1. UG", "Tiefgarage", 2400, None, None, "68 PP, Velokeller, Waschküche"],
        ["2. UG", "Technik / Keller", 1800, None, None, "HLKS-Zentrale, Kellerabteile"],
        ["Dach", "PV-Anlage", 620, None, None, "180 kWp"],
        ["Dach", "Biodiversitätsdach", 330, None, None, "Extensivbegrünung"],
    ],
    col_widths=[12, 25, 14, 12, 12, 30],
    sheet_name="Raumprogramm"
)

# Terminplan.xlsx
styled_xlsx("01 Planung/Terminplan.xlsx",
    ["Nr.", "Meilenstein", "Termin", "Dauer", "Status", "Bemerkung"],
    [
        ["M01", "Baubewilligung rechtskräftig", "20.01.2026", "-", "✓ Erledigt", ""],
        ["M02", "Abbruch Bestandsgebäude", "15.03.2026", "6 Wo.", "✓ Erledigt", ""],
        ["M03", "Baugrubenaushub fertig", "30.06.2026", "14 Wo.", "▶ Laufend", "Bohrpfähle inkl."],
        ["M04", "Rohbau UG fertig", "30.09.2026", "3 Mt.", "○ Offen", ""],
        ["M05", "Rohbau gesamt fertig", "30.04.2027", "7 Mt.", "○ Offen", ""],
        ["M06", "Fassade / Dach dicht", "30.09.2027", "5 Mt.", "○ Offen", ""],
        ["M07", "Innenausbau fertig", "30.04.2028", "7 Mt.", "○ Offen", ""],
        ["M08", "TGA-Abnahmen", "15.05.2028", "2 Wo.", "○ Offen", ""],
        ["M09", "Bauabnahme / Bezugsfreigabe", "31.05.2028", "2 Wo.", "○ Offen", ""],
        ["M10", "Bezug Wohnungen", "15.06.2028", "-", "○ Offen", ""],
    ],
    col_widths=[8, 32, 14, 10, 14, 25],
    sheet_name="Meilensteine"
)

# Architekturbriefing.docx
doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'

doc.add_heading('Gestaltungskonzept', level=0)
doc.add_paragraph('Überbauung Seefeld — Seefeldstrasse 128, 8008 Zürich')
doc.add_paragraph('')

doc.add_heading('Architektonische Leitidee', level=1)
doc.add_paragraph(
    'Der Entwurf interpretiert die für das Zürcher Seefeld-Quartier charakteristische '
    'Blockrandbebauung in zeitgenössischer Formensprache. Das Gebäude fügt sich durch '
    'seine Höhenstaffelung, die zurückhaltende Materialsprache und die Gliederung der '
    'Fassade harmonisch in das bestehende Ortsbild ein.'
)
doc.add_paragraph(
    'Der begrünte Innenhof bildet das Herzstück der Anlage. Er schafft einen geschützten '
    'Aussenraum, der Aufenthaltsqualität bietet und gleichzeitig das Mikroklima verbessert. '
    'Die Erdgeschosszone wird durch grosszügige Schaufenster zur Seefeldstrasse hin '
    'transparent gestaltet und belebt den Strassenraum.'
)

doc.add_heading('Materialkonzept', level=1)
materials = [
    ('Fassade', 'Vorfabrizierte Sichtbetonelemente mit Besenstrich-Struktur, Farbton «Sandstein hell» (RAL 085 80 20). Sockelzone in geschliffenem Ortbeton.'),
    ('Fenster', 'Holz-Metall-Fenster (Eiche / Aluminium eloxiert), 3-fach-Verglasung. Bodentiefe Fenster in Wohnräumen.'),
    ('Sonnenschutz', 'Aussenliegende textile Raffstoren (Serge Ferrari), Farbton Grau, mit Windautomatik.'),
    ('Innenhof', 'Natursteinbelag (Tessiner Gneis), heimische Bepflanzung, wassergebundene Wege.'),
    ('Wohnungen', 'Eichenparkett (Landhausdiele, geölt), Wände in Weissputz (Q3), Decken Sichtbeton geschliffen.'),
    ('Nasszellen', 'Feinsteinzeug grossformatig (60×120 cm), bodenebene Duschen, Armaturen Edelstahl.'),
]
for title, desc in materials:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_heading('Farbkonzept', level=1)
doc.add_paragraph(
    'Die Farbgebung orientiert sich an natürlichen, warmen Tönen: Sandstein und Beton '
    'an der Fassade, Eichenholz und Weissputz im Inneren. Akzente setzen die '
    'anthrazitfarbenen Metallrahmen der Fenster und Balkongeländer. Das Farbkonzept '
    'unterstützt die ruhige, hochwertige Atmosphäre des Seefeld-Quartiers.'
)

doc.save(os.path.join(DEMO, "01 Planung", "Gestaltungskonzept.docx"))
print("  ✓ 01 Planung/Gestaltungskonzept.docx")

# ══════════════════════════════════════════════════════════════════════════════
# 02 BEWILLIGUNGEN
# ══════════════════════════════════════════════════════════════════════════════

# Baubewilligung.pdf (multi-page official doc)
pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=25)

def h1(t): pdf.set_font('Helvetica','B',16); pdf.cell(0,10,t,new_x='LMARGIN',new_y='NEXT'); pdf.ln(2)
def h2(t): pdf.set_font('Helvetica','B',12); pdf.cell(0,8,t,new_x='LMARGIN',new_y='NEXT'); pdf.ln(1)
def p(t): pdf.set_font('Helvetica','',10); pdf.multi_cell(0,5.5,t); pdf.ln(2)
def bullet(t):
    pdf.set_font('Helvetica','',10)
    x = pdf.get_x()
    pdf.cell(10, 5.5, '  - ', new_x='END')
    pdf.multi_cell(pdf.w - pdf.r_margin - pdf.get_x(), 5.5, t)
    pdf.ln(1)

pdf.add_page()
pdf.set_font('Helvetica','B',11)
pdf.cell(0,6,'KANTON ZUERICH',new_x='LMARGIN',new_y='NEXT')
pdf.set_font('Helvetica','',10)
pdf.cell(0,5,'Baudirektion',new_x='LMARGIN',new_y='NEXT')
pdf.cell(0,5,'Amt fuer Raumentwicklung',new_x='LMARGIN',new_y='NEXT')
pdf.ln(15)
pdf.set_font('Helvetica','B',20)
pdf.cell(0,12,'Baubewilligungsbescheid',align='C',new_x='LMARGIN',new_y='NEXT')
pdf.ln(10)

p('Aktenzeichen: ARE-2025-14892-BB')
p('Datum: 20. Januar 2026')
p('Bauherrschaft: Seefeld Immobilien AG, Seefeldstrasse 42, 8008 Zuerich')
p('Grundstueck: Kat.-Nr. 8008-1247, Seefeldstrasse 128, 8008 Zuerich')
p('Bauvorhaben: Abbruch Bestandsgebaeude und Neubau eines siebengeschossigen Wohn- und Geschaeftshauses mit Tiefgarage')
p('Wohnzone: W5 gemaess BZO Stadt Zuerich')

pdf.ln(3)
h2('Beschluss')
p('Gestuezt auf Art. 22 RPG, $$ 218 ff. PBG und die kantonale Bauverordnung wird die Baubewilligung fuer das oben bezeichnete Vorhaben erteilt.')

pdf.add_page()
h2('Auflagen')
auflagen = [
    '1. Die Bauarbeiten sind auf Montag bis Freitag 07:00-19:00 Uhr und Samstag 08:00-17:00 Uhr zu beschraenken. Sonn- und Feiertagsarbeit ist untersagt.',
    '2. Vor Beginn der Abbruch- und Aushubarbeiten ist eine Beweissicherung der angrenzenden Liegenschaften durchzufuehren.',
    '3. Der Baumbestand gemaess Baumschutzverordnung der Stadt Zuerich ist zu erhalten. Die drei Platanen entlang der Bellerivestrasse sind waehrend der Bauzeit zu schuetzen.',
    '4. Die Anforderungen der Laermschutzverordnung (LSV) sind einzuhalten. Waehrend der Bauphase sind laermintensive Arbeiten auf ein Minimum zu beschraenken.',
    '5. Das Brandschutzkonzept ist vor Baubeginn durch die Kantonale Feuerpolizei zu genehmigen.',
    '6. Der Nachweis der Erdbebensicherheit gemaess SIA 261 ist vor Beginn der Rohbauarbeiten einzureichen.',
    '7. Die Zufahrt zur Tiefgarage ueber die Bellerivestrasse ist gemaess den Vorgaben des Tiefbauamts auszufuehren.',
    '8. Fuer die Grundwasserabsenkung waehrend der Bauphase ist eine separate Bewilligung beim AWEL einzuholen.',
    '9. Der Minergie-P-ECO-Nachweis ist spaetestens bei der Schlussabnahme vorzulegen.',
    '10. Die Photovoltaik-Anlage ist gemaess den Bestimmungen der MuKEn 2014 auszufuehren und vor Inbetriebnahme beim Elektrizitaetswerk anzumelden.',
]
for a in auflagen:
    p(a)

pdf.add_page()
h2('Rechtsmittelbelehrung')
p('Gegen diesen Bescheid kann innert 30 Tagen seit Zustellung beim Baurekursgericht des Kantons Zuerich schriftlich Rekurs erhoben werden. Die Rekursschrift muss einen Antrag und dessen Begruendung enthalten.')

pdf.ln(10)
p('Zuerich, 20. Januar 2026')
pdf.ln(8)
p('Amt fuer Raumentwicklung')
p('Abteilung Baubewilligungen')
pdf.ln(10)
pdf.cell(0,6,'_________________________',new_x='LMARGIN',new_y='NEXT')
pdf.cell(0,6,'lic. iur. Christine Lang',new_x='LMARGIN',new_y='NEXT')
pdf.cell(0,6,'Leiterin Baubewilligungen',new_x='LMARGIN',new_y='NEXT')

pdf.output(os.path.join(DEMO, "02 Bewilligungen", "Baubewilligung.pdf"))
print(f"  ✓ 02 Bewilligungen/Baubewilligung.pdf ({pdf.pages_count} pages)")

# Brandschutzkonzept.docx
doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'
doc.add_heading('Brandschutzkonzept', level=0)
doc.add_paragraph('Überbauung Seefeld, Seefeldstrasse 128, 8008 Zürich')
doc.add_paragraph('Erstellt: Brandschutzfachmann VKF, Büro Aegis Sicherheit AG')
doc.add_paragraph('Stand: Dezember 2025')
doc.add_paragraph('')

sections = [
    ('Gebäudekategorie', 'Gebäude mit mittlerer Höhe (Oberkante oberste Nutzfläche ≤ 25 m). Wohn- und Geschäftshaus, Nutzung gemischt (Wohnen, Gewerbe, Tiefgarage). Einordnung nach VKF-Brandschutzrichtlinie als Gebäude mit erhöhten Anforderungen aufgrund Mischnutzung und Tiefgarage.'),
    ('Flucht- und Rettungswege', 'Jedes Geschoss verfügt über zwei unabhängige Treppenhäuser (Treppenhaus Nord und Treppenhaus Süd) als vertikale Fluchtwege. Die Fluchtweglänge beträgt maximal 35 m von jedem Punkt der Nutzfläche. Die Treppenhäuser werden als Sicherheitstreppenhäuser mit Rauchschutz-Druckanlage (RDA) ausgeführt. Ein Feuerwehrlift (Treppenhaus Nord) ermöglicht den Feuerwehrzugang zu allen Geschossen.'),
    ('Brandabschnitte', 'Jede Wohnung bildet einen eigenen Brandabschnitt (REI 60). Die Gewerbeeinheiten im EG sind als separate Brandabschnitte mit REI 90 Abtrennung ausgeführt. Tiefgarage: eigener Brandabschnitt mit REI 90, Abschlüsse mit T30-Brandschutztüren. Die Haustechnikzentrale im 2. UG bildet einen separaten Brandabschnitt (REI 120).'),
    ('Löscheinrichtungen', 'Nass-Steigleitung in beiden Treppenhäusern mit Wandhydranten auf jedem Geschoss. Tiefgarage: Sprinkleranlage gemäss VKF-Richtlinie und SES-Richtlinie. Feuerlöscher (ABC-Pulver) in den Fluren jedes Geschosses. Externe Löscheinspeisung an der Seefeldstrasse für die Feuerwehr.'),
    ('Brandmeldeanlage', 'Vollschutz-Brandmeldeanlage mit automatischer Aufschaltung zur Einsatzzentrale Berufsfeuerwehr Zürich. Optische Rauchmelder in allen Fluren, Technikräumen und der Tiefgarage. Linienförmige Wärmemelder in der Tiefgarage zusätzlich zu Rauchmeldern. Handfeuermelder an allen Treppenhauseingängen und Fluchtwegausgängen.'),
    ('Rauch- und Wärmeabzug', 'Natürliche Rauch- und Wärmeabzugsanlage (NRWA) in den Treppenhäusern über motorisch betriebene Dachluken (je 1.5 m²). Maschinelle Entrauchung der Tiefgarage über Jetfans und Abluftkanal ins Freie. Rauchfreihaltung der Eingangshalle im EG über Überdruckbelüftung.'),
]
for title, desc in sections:
    doc.add_heading(title, level=1)
    doc.add_paragraph(desc)

doc.save(os.path.join(DEMO, "02 Bewilligungen", "Brandschutzkonzept.docx"))
print("  ✓ 02 Bewilligungen/Brandschutzkonzept.docx")

# UVB.docx
doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'
doc.add_heading('Umweltverträglichkeitsbericht (UVB)', level=0)
doc.add_paragraph('Kurzfassung — Überbauung Seefeld, Zürich')
doc.add_paragraph('Ernst Basler + Partner AG, Dezember 2025')
doc.add_paragraph('')

for title, desc in [
    ('Boden und Altlasten', 'Die Untersuchung des Baugrunds ergab keine Hinweise auf Altlasten. Der anfallende Aushub (ca. 22 000 m³) kann grösstenteils als unverschmutztes Material verwertet werden. Ca. 3 000 m³ schwach belastetes Material (Klasse B gemäss VVEA) muss auf einer bewilligten Deponie entsorgt werden.'),
    ('Grundwasser', 'Der Grundwasserspiegel liegt bei ca. 4.2 m unter Terrain. Während der Bauphase ist eine temporäre Grundwasserabsenkung erforderlich. Die Einleitung in die Kanalisation erfolgt nach Vorbehandlung (Absetz- und Neutralisationsbecken). Eine AWEL-Bewilligung ist eingeholt.'),
    ('Lärm', 'Die berechneten Lärmimmissionen des fertigen Gebäudes liegen innerhalb der Planungswerte gemäss LSV. Während der Bauphase werden die Grenzwerte für Baulärm eingehalten durch: zeitliche Beschränkung lärmintensiver Arbeiten, Einsatz lärmarmer Baumaschinen (Stufe IIIB), Schallschutzwände (3 m Höhe) zur Wohnbebauung im Norden.'),
    ('Luft', 'Die Bauphase wird durch Staubemissionen belastet. Massnahmen: Befeuchtung der Fahrwege, Reinigung der Baustellenausfahrt, Einsatz von Partikelfiltern auf allen Dieselmaschinen. Im Betrieb entstehen keine relevanten Luftschadstoff-Emissionen (keine fossilen Heizsysteme).'),
    ('Flora und Fauna', 'Auf dem Grundstück wurden keine geschützten Tier- oder Pflanzenarten nachgewiesen. Die drei Platanen an der Bellerivestrasse (Stammumfang > 80 cm, geschützt nach Baumschutzverordnung) bleiben erhalten. Als Kompensation für die Versiegelung wird der Innenhof naturnah gestaltet und ein Biodiversitätsdach realisiert.'),
    ('Ergebnis', 'Das Vorhaben ist umweltverträglich. Die Auflagen gemäss UVB sind in die Baubewilligung aufgenommen worden.'),
]:
    doc.add_heading(title, level=1)
    doc.add_paragraph(desc)

doc.save(os.path.join(DEMO, "02 Bewilligungen", "UVB_Kurzfassung.docx"))
print("  ✓ 02 Bewilligungen/UVB_Kurzfassung.docx")

# ══════════════════════════════════════════════════════════════════════════════
# 03 BAUAUSFÜHRUNG
# ══════════════════════════════════════════════════════════════════════════════

styled_xlsx("03 Bauausführung/Unternehmerverzeichnis.xlsx",
    ["Gewerk", "Unternehmer", "Ansprechperson", "Vertragssumme (CHF)", "Status", "Beginn", "Ende"],
    [
        ["Abbruch", "Eberhard Recycling AG", "Hr. Lüthi", 480000, "✓ Abgeschlossen", "01.02.2026", "15.03.2026"],
        ["Baumeister", "Marti AG Zürich", "Hr. Huber", 14200000, "▶ Beauftragt", "16.03.2026", "30.04.2027"],
        ["Spezialtiefbau", "Bauer Spezialtiefbau AG", "Fr. Rüegg", 2800000, "▶ Beauftragt", "20.03.2026", "30.06.2026"],
        ["Fassade", "Schmidlin AG", "Hr. Bachmann", 5600000, "Vergabe läuft", "01.05.2027", "30.09.2027"],
        ["Elektro", "Burkhalter Technics AG", "Hr. Sieber", 3200000, "▶ Beauftragt", "01.10.2027", "30.03.2028"],
        ["HLKS", "Amstein + Walthert AG", "Hr. Roth", 4800000, "▶ Beauftragt", "01.10.2027", "30.04.2028"],
        ["Aufzüge", "Schindler Aufzüge AG", "Fr. Meyer", 920000, "Angebot liegt vor", "01.01.2028", "30.04.2028"],
        ["Schreiner / Küchen", "Schmid Küchen AG", "Hr. Kessler", 1600000, "Vergabe läuft", "01.12.2027", "30.04.2028"],
        ["Bodenbeläge", "Belcolor AG Flooring", "Fr. Burri", 1100000, "Noch nicht vergeben", "01.01.2028", "30.04.2028"],
        ["Umgebung", "Ganz Landschaftsarchitekten", "Hr. Ganz", 680000, "Noch nicht vergeben", "01.05.2028", "15.06.2028"],
    ],
    col_widths=[18, 25, 16, 18, 16, 14, 14],
    sheet_name="Unternehmer"
)

# Baustellenordnung.pdf
pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=25)
pdf.add_page()
h1('Baustellenordnung')
pdf.set_font('Helvetica','',10)
pdf.cell(0,6,'Ueberbauung Seefeld, Seefeldstrasse 128, 8008 Zuerich',new_x='LMARGIN',new_y='NEXT')
pdf.cell(0,6,'Gueltig ab: 01.02.2026 | Bauleitung: Implenia Schweiz AG',new_x='LMARGIN',new_y='NEXT')
pdf.ln(5)

sections = [
    ('1. Arbeitszeiten', 'Montag bis Freitag: 07:00 - 19:00 Uhr\nSamstag: 08:00 - 17:00 Uhr\nSonn- und Feiertage: Arbeitsverbot\n\nLaermintensive Arbeiten (Abbruch, Rammen, Saegen): Mo-Fr 08:00-12:00 und 13:30-17:00'),
    ('2. Zugang und Sicherheit', 'Zufahrt ausschliesslich ueber Bellerivestrasse (Tor 1).\nAlle Personen melden sich am Baubuero an.\nPersoenliche Schutzausruestung ist Pflicht: Helm, Sicherheitsschuhe S3, Warnweste.\nBesucher nur in Begleitung einer autorisierten Person.\nFotografieren nur mit Genehmigung der Bauleitung.'),
    ('3. Ordnung und Entsorgung', 'Jedes Gewerk raeumt seinen Arbeitsbereich taeglich auf.\nMuelltrennung gemaess VVEA: Bauschutt, Holz, Metall, Kunststoff, Sonderabfall.\nGefahrstoffe sind im verschlossenen Gefahrstofflager aufzubewahren.\nKeine Materialablagerung ausserhalb der markierten Lagerflaechen.'),
    ('4. Umweltschutz', 'Fahrwege sind bei Trockenheit zu befeuchten.\nBaumaschinen mit Partikelfilter (Stufe IIIB) verwenden.\nOel- und Betriebsstofflagerung nur auf versiegelten Flaechen mit Auffangwanne.\nBaumschutz gemaess SN 640 561 fuer die drei Platanen.'),
    ('5. Notfall / Ansprechpartner', 'Sammelplatz: Parkplatz Bellerivestrasse 45\nErsthelfer: Roman Steiner (Polier), Tel. +41 79 445 12 88\nNaechstes Spital: Universitaetsspital Zuerich (2.8 km)\nBauleiter: Marco Keller, Tel. +41 58 474 00 00\nSicherheitsbeauftragte: Nina Frei, Tel. +41 44 720 15 30'),
]
for title, body_text in sections:
    h2(title)
    p(body_text)

pdf.output(os.path.join(DEMO, "03 Bauausführung", "Baustellenordnung.pdf"))
print(f"  ✓ 03 Bauausführung/Baustellenordnung.pdf ({pdf.pages_count} pages)")

# Bautagesbericht.docx
doc = Document()
doc.styles['Normal'].font.size = Pt(10)
doc.styles['Normal'].font.name = 'Calibri'
doc.add_heading('Bautagesbericht', level=0)
doc.add_paragraph('Projekt: Überbauung Seefeld | Datum: 15.04.2026 (Mittwoch)')
doc.add_paragraph('Wetter: Bewölkt, 12°C, kein Niederschlag | Bericht: M. Keller')
doc.add_paragraph('')

doc.add_heading('Personal', level=1)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Firma', 'Anzahl', 'Tätigkeit']): table.rows[0].cells[i].text = h
for r, row in enumerate([
    ['Marti AG', '14', 'Aushub Abschnitt C, Verbauarbeiten'],
    ['Bauer Spezialtiefbau', '6', 'Bohrpfahlherstellung Achse 5–8'],
    ['Vermessung Grunder', '2', 'Kontrollmessungen Verbau'],
    ['Implenia (Bauleitung)', '3', 'Bauleitung, Bauüberwachung'],
], 1):
    for c, val in enumerate(row): table.rows[r].cells[c].text = val

doc.add_heading('Geräte auf der Baustelle', level=1)
for g in ['1× Raupenbagger Liebherr R 946', '1× Drehbohrgerät Bauer BG 30', '3× LKW 4-Achser (Erdtransport)', '1× Pneukran Liebherr LTM 1100', '1× Radlader CAT 950']:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('Ausgeführte Arbeiten', level=1)
doc.add_paragraph('Aushub Abschnitt C auf Kote -5.80 m (Tagesziel -6.20 m nicht erreicht wegen Felslinse).')
doc.add_paragraph('Bohrpfähle Nr. 24–26 hergestellt (Ø 80 cm, Tiefe 18.5 m). Betonierprotokoll liegt vor.')
doc.add_paragraph('Spritzbeton-Sicherung Verbau Achse 3–5 abgeschlossen.')
doc.add_paragraph('Ca. 340 m³ Aushub abgefahren (Deponie Ebikofer, Dietlikon).')

doc.add_heading('Besondere Vorkommnisse', level=1)
doc.add_paragraph('Felslinse in Abschnitt C angetroffen (ca. 3×4 m). Felsabbau mit hydraulischem Meissel erforderlich. Voraussichtlicher Mehraufwand: 2 Arbeitstage.')

doc.add_heading('Pendenzen / Nächste Schritte', level=1)
for s in ['Felsabbau Abschnitt C fortsetzen', 'Bohrpfähle Nr. 27–30 herstellen', 'Grundwassermessung (2× wöchentlich)', 'Beweissicherung Nachbarliegenschaft Nr. 130 (Rissprotokoll)']:
    doc.add_paragraph(s, style='List Bullet')

doc.save(os.path.join(DEMO, "03 Bauausführung", "Bautagesbericht_2026-04-15.docx"))
print("  ✓ 03 Bauausführung/Bautagesbericht_2026-04-15.docx")

# Mängelliste.xlsx
styled_xlsx("03 Bauausführung/Mängelliste.xlsx",
    ["Nr.", "Gewerk", "Beschreibung", "Ort", "Festgestellt", "Frist", "Status", "Verantwortlich"],
    [
        ["M-001", "Spezialtiefbau", "Bohrpfahl Nr. 12: Betonüberverbrauch 18% (Hohlraumbildung)", "Achse 3, 1. UG", "02.04.2026", "15.04.2026", "In Prüfung", "Bauer Spezialtiefbau"],
        ["M-002", "Baumeister", "Verbau Achse 6: Horizontalverschiebung 4 cm (Toleranz 3 cm)", "Nordseite", "08.04.2026", "22.04.2026", "Offen", "Marti AG"],
        ["M-003", "Baumeister", "Entwässerungsrinne Baustellenzufahrt verstopft", "Tor 1", "10.04.2026", "12.04.2026", "✓ Behoben", "Marti AG"],
    ],
    col_widths=[8, 16, 40, 16, 14, 14, 12, 20],
    sheet_name="Mängel"
)

# ══════════════════════════════════════════════════════════════════════════════
# 04 KOSTEN
# ══════════════════════════════════════════════════════════════════════════════

styled_xlsx("04 Kosten/Kostenvoranschlag_BKP.xlsx",
    ["BKP", "Bezeichnung", "Kosten (CHF)", "Anteil (%)", "Bemerkung"],
    [
        [1, "Vorbereitungsarbeiten", 1850000, 2.9, "Abbruch, Baustelleneinrichtung"],
        [2, "Gebäude", 42600000, 67.8, ""],
        ["  21", "  Rohbau 1", 12400000, None, "Baumeister, Spezialtiefbau"],
        ["  22", "  Rohbau 2", 3200000, None, "Montagebau, vorfabrizierte Elemente"],
        ["  23", "  Elektroanlagen", 3200000, None, "Stark-/Schwachstrom, PV-Anlage"],
        ["  24", "  HLK-Anlagen", 4100000, None, "Heizung, Lüftung, Klima, Kälte"],
        ["  25", "  Sanitäranlagen", 2400000, None, "Sanitär, Sprinkler"],
        ["  26", "  Transportanlagen", 920000, None, "2 Aufzüge"],
        ["  27", "  Ausbau 1", 8600000, None, "Fenster, Fassade, Dach"],
        ["  28", "  Ausbau 2", 7780000, None, "Bodenbeläge, Schreiner, Maler"],
        [3, "Betriebseinrichtungen", 1200000, 1.9, "Küchen, Waschküche"],
        [4, "Umgebung", 2400000, 3.8, "Umgebung, Tiefgarage Zufahrt"],
        [5, "Baunebenkosten", 6250000, 10.0, "Honorare, Gebühren, Gutachten"],
        [6, "Reserve", 3500000, 5.6, "Unvorhergesehenes, Teuerung"],
        [9, "Ausstattung", 5000000, 8.0, "Möblierung Gewerbe, Allgemeinräume"],
        [None, "TOTAL BKP 1–9", 62800000, 100.0, ""],
    ],
    col_widths=[10, 35, 18, 12, 32],
    sheet_name="BKP Kosten"
)

styled_xlsx("04 Kosten/Zahlungsplan_2026.xlsx",
    ["Monat", "Geplant (CHF)", "Kumuliert (CHF)", "Ist (CHF)", "Abweichung (CHF)", "Bemerkung"],
    [
        ["Feb 2026", 320000, 320000, 335000, 15000, "Abbruch, Baustelleneinrichtung"],
        ["Mär 2026", 580000, 900000, 560000, -20000, ""],
        ["Apr 2026", 1200000, 2100000, 1180000, -20000, "Spezialtiefbau begonnen"],
        ["Mai 2026", 1600000, 3700000, None, None, "Prognose"],
        ["Jun 2026", 1800000, 5500000, None, None, "Prognose"],
        ["Jul 2026", 2200000, 7700000, None, None, "Rohbau startet"],
        ["Aug 2026", 2400000, 10100000, None, None, "Prognose"],
        ["Sep 2026", 2200000, 12300000, None, None, "Prognose"],
        ["Okt 2026", 2000000, 14300000, None, None, "Prognose"],
        ["Nov 2026", 1800000, 16100000, None, None, "Prognose"],
        ["Dez 2026", 1200000, 17300000, None, None, "Winterpause, reduziert"],
    ],
    col_widths=[14, 16, 16, 16, 16, 28],
    sheet_name="Zahlungsplan"
)

# Nachtragsverzeichnis.docx
doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'
doc.add_heading('Nachtragsverzeichnis', level=0)
doc.add_paragraph('Überbauung Seefeld — Stand April 2026')
doc.add_paragraph('')

for nt, data in [
    ('NT-001: Mehrkosten Felsabbau', {
        'Unternehmer': 'Marti AG Zürich',
        'Forderung': 'CHF 125 000.– netto',
        'Begründung': 'Unerwartete Felslinsen in den Abschnitten B und C. Das Baugrundgutachten (Emch+Berger, Juli 2025) hatte homogenen Kies/Sand prognostiziert. Hydraulischer Felsabbau mit Meissel und teilweise Sprengvortrieb (Kleinsprengungen) erforderlich.',
        'Bewertung BL': 'Grundsätzlich berechtigt. Empfehlung: CHF 95 000.– anerkennen (Meissel-Arbeiten nachvollziehbar, Sprengkosten teilweise überhöht).',
        'Status': '🔄 In Prüfung durch Bauherrschaft',
    }),
    ('NT-002: Zusätzliche Grundwasserhaltung', {
        'Unternehmer': 'Bauer Spezialtiefbau AG',
        'Forderung': 'CHF 48 000.– netto',
        'Begründung': 'Grundwasserzufluss ca. 30% höher als im hydrogeologischen Gutachten prognostiziert. Zusätzliche Pumpenkapazität und verlängerte Laufzeit (2 Wochen) erforderlich.',
        'Bewertung BL': 'Berechtigt. Mengenmehrung nachvollziehbar dokumentiert.',
        'Status': '✓ Genehmigt durch Bauherrschaft (18.04.2026)',
    }),
]:
    doc.add_heading(nt, level=1)
    for k, v in data.items():
        doc.add_paragraph(f'{k}: {v}')
    doc.add_paragraph('')

doc.add_paragraph('Summe Nachtragsforderungen: CHF 173 000.– netto')
doc.add_paragraph('Davon genehmigt: CHF 48 000.–')
doc.add_paragraph('Davon in Prüfung: CHF 125 000.–')

doc.save(os.path.join(DEMO, "04 Kosten", "Nachtragsverzeichnis.docx"))
print("  ✓ 04 Kosten/Nachtragsverzeichnis.docx")

# ══════════════════════════════════════════════════════════════════════════════
# 05 PROTOKOLLE
# ══════════════════════════════════════════════════════════════════════════════

# Multi-page meeting minutes
doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'
doc.add_heading('Bausitzung Nr. 14', level=0)
doc.add_paragraph('Datum: 16.04.2026, 09:30–11:15 Uhr')
doc.add_paragraph('Ort: Baubüro, Seefeldstrasse 128, Zürich')
doc.add_paragraph('')

doc.add_heading('Teilnehmende', level=1)
for p_name in [
    'Dr. Stefan Meier (Bauherrschaft)',
    'Arch. ETH Laura Brunner (Generalplanung)',
    'Dipl. Bauing. ETH Marco Keller (Bauleitung)',
    'Ing. Nina Frei (SIBE)',
    'Peter Huber (Marti AG)',
    'Anna Rüegg (Bauer Spezialtiefbau)',
    'Dr. Ing. Thomas Wyss (Tragwerk)',
]:
    doc.add_paragraph(p_name, style='List Bullet')

doc.add_heading('Traktandum 1: Terminübersicht', level=1)
doc.add_paragraph(
    'Der Aushub liegt ca. 4 Arbeitstage hinter dem Plan. Ursache: Felslinsen in den '
    'Abschnitten B und C (siehe Nachtrag NT-001). Die Bauleitung hat einen zusätzlichen '
    'Hydraulikbagger organisiert, der ab KW 17 im Einsatz ist.'
)
doc.add_paragraph('Prognose: Rückstand kann bis Ende Mai aufgeholt werden.')
doc.add_paragraph('Verantwortlich: Marti AG / Bauleitung | Termin: 30.05.2026')

doc.add_heading('Traktandum 2: Spezialtiefbau', level=1)
doc.add_paragraph(
    'Bauer Spezialtiefbau meldet: 26 von 42 Bohrpfählen sind hergestellt. Qualität '
    'der Pfähle ist einwandfrei (Integritätsprüfung bei 8 Pfählen durchgeführt, alle i.O.). '
    'Bei Pfahl Nr. 12 wurde ein Betonüberverbrauch von 18% festgestellt (Hohlraumbildung '
    'im Lockergestein). Statische Unbedenklichkeit wird durch Tragwerksplaner bestätigt.'
)
doc.add_paragraph('Fertigstellung aller Bohrpfähle: voraussichtlich 15.05.2026.')

doc.add_heading('Traktandum 3: Kosten und Nachträge', level=1)
doc.add_paragraph(
    'Nachtrag NT-001 (Felsabbau, CHF 125 000.–) wird durch Bauleitung geprüft. '
    'Empfehlung liegt bis zur nächsten Sitzung vor. Nachtrag NT-002 (Grundwasserhaltung, '
    'CHF 48 000.–) wurde in der Sitzung durch die Bauherrschaft genehmigt.'
)
doc.add_paragraph('Aktueller Budgetstatus: CHF 2.075 Mio. ausgegeben von CHF 62.8 Mio. (3.3%)')

doc.add_heading('Traktandum 4: Arbeitssicherheit', level=1)
doc.add_paragraph(
    'SIBE-Bericht (N. Frei): Begehung am 12.04.2026. Keine wesentlichen Beanstandungen. '
    'Ein Verstoss gegen die Helmtragepflicht wurde mündlich gerügt. Die Absturzsicherung '
    'an der Baugrubenoberkante ist vollständig installiert.'
)
doc.add_paragraph('Nächste SIBE-Begehung: 26.04.2026')

doc.add_heading('Traktandum 5: Verschiedenes', level=1)
doc.add_paragraph('Bauherrschaft wünscht monatliche Drohnen-Dokumentation (Beginn Mai 2026).')
doc.add_paragraph('Nächste Bausitzung: 23.04.2026, 09:30 Uhr')
doc.add_paragraph('')
doc.add_paragraph('Protokoll: M. Keller, 16.04.2026').italic = True

doc.save(os.path.join(DEMO, "05 Protokolle", "Bausitzung_Nr14_2026-04-16.docx"))
print("  ✓ 05 Protokolle/Bausitzung_Nr14_2026-04-16.docx")

# Baufortschrittsbericht (long PDF)
pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=25)

# Cover
pdf.add_page()
pdf.set_font('Helvetica','B',28)
pdf.ln(35)
pdf.cell(0,15,'Baufortschrittsbericht',align='C',new_x='LMARGIN',new_y='NEXT')
pdf.set_font('Helvetica','',16)
pdf.cell(0,10,'Ueberbauung Seefeld, Zuerich',align='C',new_x='LMARGIN',new_y='NEXT')
pdf.ln(8)
pdf.set_font('Helvetica','',12)
pdf.cell(0,8,'Berichtszeitraum: Maerz / April 2026',align='C',new_x='LMARGIN',new_y='NEXT')
pdf.cell(0,8,'Implenia Schweiz AG | Dipl. Bauing. ETH Marco Keller',align='C',new_x='LMARGIN',new_y='NEXT')
pdf.ln(40)
pdf.set_font('Helvetica','',9)
pdf.cell(0,5,'Vertraulich - Nur fuer Projektbeteiligte',align='C',new_x='LMARGIN',new_y='NEXT')

# TOC
pdf.add_page()
h1('Inhaltsverzeichnis')
for item in ['1. Zusammenfassung','2. Terminuebersicht','3. Baufortschritt im Detail','4. Kostenuebersicht','5. Qualitaet und Maengel','6. Arbeitssicherheit','7. Nachtraege','8. Risiken und Massnahmen','9. Ausblick Mai 2026','10. Fotodokumentation (Verweis)']:
    pdf.set_font('Helvetica','',11)
    pdf.cell(0,7,item,new_x='LMARGIN',new_y='NEXT')

# Content pages
pdf.add_page()
h1('1. Zusammenfassung')
p('Der Baufortschritt in den Monaten Maerz und April 2026 verlief weitgehend planmaessig. Der Abbruch des Bestandsgebaeudes wurde termingerecht abgeschlossen. Die Baugrube ist zu ca. 65% ausgehoben. Die Bohrpfahlarbeiten sind zu ca. 60% abgeschlossen (26 von 42 Pfaehlen).')
p('Wesentliche Ereignisse:')
for e in ['Abbruch Bestandsgebaeude abgeschlossen (15.03.2026)','Spezialtiefbau (Bohrpfaehle) begonnen am 20.03.2026','Felslinsen in Abschnitten B/C entdeckt - Nachtrag NT-001','Grundwasserzufluss hoeher als erwartet - Nachtrag NT-002 genehmigt','Keine Arbeitsunfaelle im Berichtszeitraum']:
    bullet(e)
pdf.ln(3)
p('Gesamtbewertung: GELB (bedingt planmaessig). Ursache ist eine Verzoegerung von ca. 4 Arbeitstagen durch die Felslinsen. Gegenmassnahmen sind eingeleitet.')

pdf.add_page()
h1('2. Terminuebersicht')
p('Aktueller Stand der Meilensteine:')
pdf.set_font('Helvetica','B',9)
cw = (pdf.w-40)/4
for hdr in ['Nr.','Meilenstein','Termin','Status']:
    pdf.cell(cw,7,hdr,border=1,align='C')
pdf.ln()
pdf.set_font('Helvetica','',9)
for row in [['M01','Baubewilligung','20.01.2026','Erledigt'],['M02','Abbruch','15.03.2026','Erledigt'],['M03','Baugrube fertig','30.06.2026','Laufend'],['M04','Rohbau UG','30.09.2026','Offen'],['M05','Rohbau gesamt','30.04.2027','Offen']]:
    for c in row: pdf.cell(cw,6,c,border=1)
    pdf.ln()
pdf.ln(5)
h2('Terminabweichungen')
p('Die Erdarbeiten liegen aktuell ca. 4 Arbeitstage hinter dem Plan. Ursache sind Felslinsen im Untergrund, die mit hydraulischem Meissel abgebaut werden muessen. Ab KW 17 ist ein zusaetzlicher Bagger im Einsatz. Prognose: Rueckstand kann bis Ende Mai 2026 aufgeholt werden.')

pdf.add_page()
h1('3. Baufortschritt im Detail')
h2('3.1 Abbrucharbeiten')
p('Die Abbrucharbeiten wurden von der Eberhard Recycling AG termingerecht durchgefuehrt und am 15.03.2026 abgeschlossen. Insgesamt wurden ca. 4200 m3 Abbruchmaterial entsorgt, davon 78% recycelt (Betonrecycling). Die Altlastenuntersuchung ergab keine Belastungen.')
h2('3.2 Baugrube und Verbau')
p('Der Aushub der Baugrube ist zu ca. 65% abgeschlossen. Die Verbauarbeiten (rueckverankerte Bohrpfahlwand) verlaufen planmaessig. Die Inklinometermessungen zeigen Verformungen innerhalb der zulaessigen Grenzwerte.')
p('Aushubvolumen bisher: ca. 14500 m3 von geschaetzten 22000 m3.')
h2('3.3 Spezialtiefbau')
p('26 von 42 Bohrpfaehlen sind hergestellt (62%). Die Integritaetspruefungen (Low-Strain-Test) an 8 Pfaehlen waren alle ohne Befund. Bei Pfahl Nr. 12 wurde ein Betonueberverbrauch von 18% festgestellt (Hohlraumbildung im Lockergestein). Die statische Unbedenklichkeit wurde durch den Tragwerksplaner bestaetigt.')

pdf.add_page()
h1('4. Kostenuebersicht')
p('Kostenstand per 30.04.2026:')
pdf.set_font('Helvetica','B',9)
for hdr in ['Position','Budget (CHF)','Ist (CHF)','Prognose (CHF)']:
    pdf.cell(cw,7,hdr,border=1,align='C')
pdf.ln()
pdf.set_font('Helvetica','',9)
for row in [['Abbruch','480\'000','480\'000','480\'000'],['Baumeister','14\'200\'000','1\'180\'000','14\'350\'000'],['Spezialtiefbau','2\'800\'000','840\'000','2\'848\'000'],['Baustelleneinr.','350\'000','335\'000','345\'000'],['Summe bisher','17\'830\'000','2\'835\'000','18\'023\'000']]:
    for c in row: pdf.cell(cw,6,c,border=1)
    pdf.ln()
pdf.ln(5)
p('Die Gesamtkosten liegen aktuell 1.1% ueber dem Budget. Die Ueberschreitung resultiert aus den Mehrkosten fuer Felsabbau (NT-001, noch in Pruefung) und der hoeheren Grundwasserhaltung (NT-002, CHF 48\'000 genehmigt).')

pdf.add_page()
h1('5. Qualitaet und Maengel')
p('Durchgefuehrte Qualitaetspruefungen:')
for c in ['Verdichtungspruefung Baugrubensohle Abschnitt A: Bestanden','Integritaetspruefung Bohrpfaehle Nr. 5, 8, 12, 15, 18, 20, 23, 26: Alle i.O.','Betonpruefung: 6 Wuerfel entnommen, 28-Tage-Pruefung ausstehend','Inklinometermessungen: Alle Verformungen im Toleranzbereich']:
    bullet(c)
pdf.ln(3)
h2('Offene Maengel')
p('M-001: Bohrpfahl Nr. 12 Betonueberverbrauch - In Pruefung')
p('M-002: Verbau Achse 6 Horizontalverschiebung - Offen')
p('M-003: Entwässerungsrinne Zufahrt - Behoben')

h1('6. Arbeitssicherheit')
p('Im Berichtszeitraum wurden keine Arbeitsunfaelle registriert. 4 SIBE-Begehungen durchgefuehrt. 18 Erstunterweisungen. 1 Verstoss gegen Helmtragepflicht (muendliche Ermahnung).')

pdf.add_page()
h1('8. Risiken und Massnahmen')
pdf.set_font('Helvetica','B',9)
for hdr in ['Risiko','Wahrsch.','Auswirkung','Massnahme']:
    pdf.cell(cw,7,hdr,border=1,align='C')
pdf.ln()
pdf.set_font('Helvetica','',9)
for row in [['Weitere Felslinsen','Mittel','Terminverzug','Meissel bereit'],['Grundwasser','Gering','Mehrkosten','Zusatzpumpen'],['Lieferengpaesse','Gering','Terminverzug','Fruehbestellung'],['Laermklagen','Mittel','Bauverzoeg.','Schutzwand erhoehen']]:
    for c in row: pdf.cell(cw,6,c,border=1)
    pdf.ln()
pdf.ln(5)

h1('9. Ausblick Mai 2026')
for o in ['Fertigstellung Aushub Abschnitt C und D','Restliche 16 Bohrpfaehle herstellen','Beginn Sauberkeitsschicht und Bodenplatte (Abschnitt A)','Erste Drohnen-Dokumentation','Naechste Bausitzung: 07.05.2026']:
    bullet(o)
pdf.ln(5)
p('Zuerich, den 30.04.2026')
pdf.ln(8)
pdf.cell(0,6,'_________________________',new_x='LMARGIN',new_y='NEXT')
pdf.cell(0,6,'Dipl. Bauing. ETH Marco Keller',new_x='LMARGIN',new_y='NEXT')
pdf.cell(0,6,'Implenia Schweiz AG',new_x='LMARGIN',new_y='NEXT')

pdf.output(os.path.join(DEMO, "05 Protokolle", "Baufortschrittsbericht_Apr_2026.pdf"))
print(f"  ✓ 05 Protokolle/Baufortschrittsbericht_Apr_2026.pdf ({pdf.pages_count} pages)")

# ══════════════════════════════════════════════════════════════════════════════
# 06 FOTOS (keep existing JPGs, just rename)
# ══════════════════════════════════════════════════════════════════════════════
# Fotos are already real JPGs from Unsplash — we keep them as-is
print("  ✓ 06 Fotos/ (keeping existing JPG files)")

# ══════════════════════════════════════════════════════════════════════════════
# 07 VERTRÄGE
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'
doc.add_heading('Generalplanervertrag (Entwurf)', level=0)
doc.add_paragraph('')

doc.add_paragraph('zwischen')
doc.add_paragraph('Seefeld Immobilien AG, Seefeldstrasse 42, 8008 Zürich')
doc.add_paragraph('(nachfolgend «Auftraggeberin»)')
doc.add_paragraph('')
doc.add_paragraph('und')
doc.add_paragraph('Baumschlager Eberle Architekten AG, Limmatstrasse 65, 8005 Zürich')
doc.add_paragraph('(nachfolgend «Auftragnehmerin»)')
doc.add_paragraph('')

for title, text in [
    ('Art. 1 Gegenstand', 'Die Auftragnehmerin erbringt die Generalplanerleistungen für den Neubau eines Wohn- und Geschäftshauses auf der Liegenschaft Seefeldstrasse 128, 8008 Zürich, gemäss den Leistungsphasen 1–6 nach SIA 112 (Modell Bauplanung).'),
    ('Art. 2 Leistungsumfang', 'Die Generalplanerin koordiniert folgende Fachplanungen:\n– Architektur (Eigenleistung)\n– Tragwerksplanung (Schnetzer Puskas Ingenieure)\n– HLKS-Planung (Amstein + Walthert AG)\n– Elektroplanung (Burkhalter Technics AG)\n– Bauphysik (Gartenmann Engineering)\n– Brandschutz (Aegis Sicherheit AG)\n– Landschaftsarchitektur (Ganz Landschaftsarchitekten)'),
    ('Art. 3 Honorar', 'Das Honorar wird auf Grundlage der SIA-Honorarordnung 102/103 berechnet:\n– Auftragsvolumen (BKP 2): CHF 42.6 Mio.\n– Honorarsatz Generalplanung: 12.8%\n– Geschätztes Gesamthonorar: ca. CHF 5.45 Mio. (exkl. MwSt.)\n\nDie Abrechnung erfolgt nach Leistungsfortschritt gemäss SIA 112.'),
    ('Art. 4 Termine', 'Die Generalplanerin verpflichtet sich, die Planungsleistungen gemäss dem vereinbarten Terminplan zu erbringen. Wesentliche Planungsfristen:\n– Bauprojekt (Phase 31): abgeschlossen\n– Bewilligungsplanung (Phase 32): abgeschlossen\n– Ausführungsplanung (Phase 41): laufend, Abschluss bis 30.06.2026\n– Ausschreibung (Phase 51): rollierend nach Gewerk'),
    ('Art. 5 Haftung und Versicherung', 'Die Auftragnehmerin haftet für die sorgfältige Erbringung ihrer Leistungen gemäss den anerkannten Regeln der Baukunde und den SIA-Normen. Sie unterhält eine Berufshaftpflichtversicherung mit einer Deckungssumme von mindestens CHF 10 Mio. pro Schadensfall.'),
    ('Art. 6 Gerichtsstand', 'Für Streitigkeiten aus diesem Vertrag sind die Gerichte am Sitz der Auftraggeberin (Zürich) zuständig. Es gilt Schweizer Recht.'),
]:
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)

doc.save(os.path.join(DEMO, "07 Verträge", "Generalplanervertrag.docx"))
print("  ✓ 07 Verträge/Generalplanervertrag.docx")

styled_xlsx("07 Verträge/Bürgschaftsübersicht.xlsx",
    ["Firma", "Art", "Betrag (CHF)", "Bank / Versicherer", "Gültig bis", "Status"],
    [
        ["Marti AG Zürich", "Erfüllungsgarantie", 1420000, "Zürcher Kantonalbank", "30.04.2028", "✓ Vorliegend"],
        ["Marti AG Zürich", "Gewährleistungsgarantie", 710000, "Zürcher Kantonalbank", "30.04.2033", "Noch nicht fällig"],
        ["Bauer Spezialtiefbau", "Erfüllungsgarantie", 280000, "Credit Suisse", "30.06.2027", "✓ Vorliegend"],
        ["Schmidlin AG", "Erfüllungsgarantie", 560000, "Basler Versicherung", "30.09.2028", "Ausstehend"],
        ["Burkhalter Technics", "Erfüllungsgarantie", 320000, "Zürich Versicherung", "30.03.2029", "✓ Vorliegend"],
    ],
    col_widths=[22, 22, 16, 22, 14, 16],
    sheet_name="Bürgschaften"
)

# ══════════════════════════════════════════════════════════════════════════════
# 08 SIBE
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.name = 'Calibri'
doc.add_heading('Sicherheitskonzept (SiKo)', level=0)
doc.add_paragraph('Überbauung Seefeld, Seefeldstrasse 128, 8008 Zürich')
doc.add_paragraph('SIBE: Ing. Nina Frei, SIBE Zürich GmbH | Stand: April 2026')
doc.add_paragraph('')

doc.add_heading('1. Gefährdungsbeurteilung — Aktuelle Bauphase', level=1)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['Gefährdung', 'Massnahme', 'Verantwortlich']): table.rows[0].cells[i].text = h
for r, row in enumerate([
    ['Absturz in Baugrube', 'Geländer (h=1.10 m), Zugangsleiter, Beleuchtung', 'Polier'],
    ['Baugrubeneinbruch', 'Verbau gemäss Statik, tägliche Sichtprüfung', 'Bauleitung'],
    ['Kollision Maschinen/Personen', 'Getrennte Verkehrswege, Einweiser bei Rückfahrt', 'Polier'],
    ['Lärm (> 85 dB)', 'Gehörschutz, lärmreduzierte Bauverfahren', 'Alle Gewerke'],
    ['Grundwasserkontakt', 'Wasserdichte Schutzkleidung, Hautschutzplan', 'Marti AG'],
    ['Herabfallende Gegenstände', 'Fangnetze, Helmpflicht, Absperrung', 'Polier'],
], 1):
    for c, val in enumerate(row): table.rows[r].cells[c].text = val

doc.add_heading('2. Persönliche Schutzausrüstung (PSA)', level=1)
for item in ['Schutzhelm EN 397', 'Sicherheitsschuhe S3 (EN ISO 20345)', 'Warnweste Klasse 2 (EN ISO 20471)', 'Schutzbrille bei Schleif-/Bohrarbeiten', 'Gehörschutz bei Lärmpegel > 85 dB(A)', 'Staubmaske FFP2 bei Abbruch-/Bohrarbeiten']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3. Notfallorganisation', level=1)
doc.add_paragraph('Notruf: 144 (Sanität) / 118 (Feuerwehr) / 117 (Polizei)')
doc.add_paragraph('Ersthelfer: Roman Steiner (Polier), Peter Huber (Marti AG)')
doc.add_paragraph('Erste-Hilfe-Material: Baubüro + jede Bauebene')
doc.add_paragraph('Sammelplatz: Parkplatz Bellerivestrasse 45')
doc.add_paragraph('Nächstes Spital: Universitätsspital Zürich (2.8 km / 8 Min.)')

doc.save(os.path.join(DEMO, "08 SIBE", "Sicherheitskonzept.docx"))
print("  ✓ 08 SIBE/Sicherheitskonzept.docx")

styled_xlsx("08 SIBE/Unterweisungsnachweis_2026.xlsx",
    ["Datum", "Name", "Firma", "Thema", "Instruiert durch", "Unterschrift"],
    [
        ["01.02.2026", "Steiner, Roman", "Implenia", "Erstinstruktion Baustellenordnung", "Frei, Nina", "✓"],
        ["01.02.2026", "Huber, Peter", "Marti AG", "Erstinstruktion Baustellenordnung", "Frei, Nina", "✓"],
        ["01.02.2026", "Müller, Hans", "Eberhard Recycling", "Erstinstruktion + Abbrucharbeiten", "Frei, Nina", "✓"],
        ["15.02.2026", "Keller, Marco", "Implenia", "Erstinstruktion Baustellenordnung", "Frei, Nina", "✓"],
        ["20.03.2026", "Rüegg, Anna", "Bauer Spezialtiefbau", "Erstinstruktion + Bohrpfahlarbeiten", "Frei, Nina", "✓"],
        ["20.03.2026", "Sutter, Thomas", "Bauer Spezialtiefbau", "Erstinstruktion + Kran-/Hebearbeiten", "Frei, Nina", "✓"],
        ["01.04.2026", "Alle Arbeitnehmer", "Alle", "Sicherheitsbegehung Baugrube", "Frei, Nina", "✓"],
        ["08.04.2026", "Marti-Crew (14 Pers.)", "Marti AG", "Arbeiten in der Tiefe / Verbau", "Steiner, Roman", "✓"],
        ["15.04.2026", "Alle Arbeitnehmer", "Alle", "Verhalten bei Grundwassereintritt", "Frei, Nina", "✓"],
    ],
    col_widths=[14, 22, 20, 32, 16, 12],
    sheet_name="Unterweisungen"
)

print("\n✅ All demo files generated successfully!")
print(f"   Location: {DEMO}")
